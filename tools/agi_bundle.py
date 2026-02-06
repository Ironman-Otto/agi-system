
# ==================================================
# FILE: C:\dev\agi-system\src\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\adapters\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\adapters\llm\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\adapters\llm\llm_adapter_http.py
# ==================================================

import json
import uuid
import requests
from typing import Any

from src.core.adapters.llm.llm_prompt import INTENT_EXTRACTION_PROMPT


class HTTPLLMAdapter:
    """
    Adapter for a locally hosted LLM over HTTP.
    """

    def __init__(self, endpoint_url: str):
        self.endpoint_url = endpoint_url

    def classify_directive(self, directive_text: str) -> dict[str, Any]:
        prompt = INTENT_EXTRACTION_PROMPT.format(directive=directive_text)

        payload = {
            "prompt": prompt,
            "temperature": 0.0,
        }

        response = requests.post(self.endpoint_url, json=payload, timeout=30)
        response.raise_for_status()

        raw = response.json()["text"]
        data = json.loads(raw)
        data.setdefault("intent_id", str(uuid.uuid4()))
        return data

# ==================================================
# FILE: C:\dev\agi-system\src\core\adapters\llm\llm_adapter_openai.py
# ==================================================

import json
import uuid
from typing import Any

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

from src.core.adapters.llm.llm_prompt import INTENT_EXTRACTION_PROMPT


class OpenAILLMAdapter:
    """
    Real LLM adapter using OpenAI API.
    Safe to import even if OpenAI is not installed.
    """

    def __init__(self, model: str = "gpt-4.1-mini"):
        if OpenAI is None:
            raise RuntimeError("openai package not installed.")
        self.client = OpenAI()
        self.model = model

    def classify_directive(self, directive_text: str) -> dict[str, Any]:
        prompt = INTENT_EXTRACTION_PROMPT.format(directive=directive_text)

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": "Return JSON only."},
                {"role": "user", "content": prompt},
            ],
        )

        raw = response.choices[0].message.content
        data = json.loads(raw)

        # Ensure intent_id exists even if model forgot
        data.setdefault("intent_id", str(uuid.uuid4()))
        return data

# ==================================================
# FILE: C:\dev\agi-system\src\core\adapters\llm\llm_prompt.py
# ==================================================

INTENT_EXTRACTION_PROMPT = """
You are an intent extraction module.

Your task is to classify a directive and output a JSON object ONLY.
Do not include explanations or formatting.

Allowed directive_type values:
- cognitive
- analytical
- goal_oriented
- behavioral
- supervisory

Rules:
- planning_required is true ONLY if multi-step execution, coordination,
  or state change is required.
- confidence_score must be between 0.0 and 1.0

Required JSON fields:
intent_id
directive_source
directive_type
planning_required
urgency_level
risk_level
expected_response_type
confidence_score

Directive:
{directive}
"""

# ==================================================
# FILE: C:\dev\agi-system\src\core\agent\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\agent\agent_loop.py
# ==================================================

class AgentLoop:
    """
    The Agent Loop is the executive control structure.
    It owns the think → decide → act cycle.
    """

    def __init__(self, intent_extractor, router, behavior_registry,
                 llm_consultant=None, skill_executor=None):
        self.intent_extractor = intent_extractor
        self.router = router
        self.behavior_registry = behavior_registry
        self.llm = llm_consultant
        self.skill_executor = skill_executor
    def run(self, directive: str):
        print("\n[AgentLoop] Received directive:", directive)

        # 1. Interpret
        intent = self.intent_extractor.extract_intent(directive)

        # 2. Decide
        route = self.router.route(intent)

        if route == "direct_response":
            return {
                "status": "direct_response",
                "intent": intent
            }

        if route == "request_clarification":
            return {
                "status": "clarification_required",
                "intent": intent
            }

        if route == "invoke_planner":
            return self._handle_goal(intent)

        return {"status": "unknown_route"}

    def _handle_goal(self, intent):
        if self.llm is None:
            return {
                "status": "planner_needed",
                "intent": intent
            }

        advice = self.llm.consult(intent, self.behavior_registry)

        if advice["recommendation_type"] != "invoke_skill":
            return {
                "status": "no_action",
                "intent": intent
            }

        skill_name = advice["skill"]

        if not self.behavior_registry.is_allowed(skill_name):
            return {
                "status": "skill_not_allowed",
                "skill": skill_name
            }

        behavior = self.behavior_registry.get(skill_name)

        if behavior.requires_approval:
            return {
                "status": "approval_required",
                "skill": skill_name
            }

        # EXECUTE SKILL
        result = self.skill_executor.execute(
            skill_name,
            advice["arguments"]
        )

        return {
            "status": "skill_executed",
            "skill": skill_name,
            "result": result
        }

# ==================================================
# FILE: C:\dev\agi-system\src\core\agent\behavior_registry.py
# ==================================================

from dataclasses import dataclass
from typing import Dict, Any


@dataclass
class Behavior:
    name: str
    risk: str
    requires_approval: bool


class BehaviorRegistry:
    """
    Registry of allowed behaviors (skills).
    This is the executable surface of the agent.
    """

    def __init__(self):
        self._behaviors: Dict[str, Behavior] = {}

    def register(self, name: str, risk: str = "low", requires_approval: bool = False):
        self._behaviors[name] = Behavior(
            name=name,
            risk=risk,
            requires_approval=requires_approval
        )

    def is_allowed(self, name: str) -> bool:
        return name in self._behaviors

    def get(self, name: str) -> Behavior | None:
        return self._behaviors.get(name)

    def list_behaviors(self) -> list[str]:
        return list(self._behaviors.keys())

# ==================================================
# FILE: C:\dev\agi-system\src\core\agent\llm_consultant_stub.py
# ==================================================

class LLMConsultantStub:
    """
    Stand-in for an LLM-based cognitive advisor.
    Returns structured recommendations only.
    """

    def consult(self, intent, behavior_registry):
        # Very naive logic for now — intentional
        if intent.directive_type.value == "goal_oriented":
            if "create_docx" in behavior_registry.list_behaviors():
                return {
                    "recommendation_type": "invoke_skill",
                    "skill": "create_docx",
                    "arguments": {
                        "title": "Generated Plan",
                        "sections": ["Overview", "Steps", "Validation"]
                    },
                    "confidence": 0.7
                }

        return {
            "recommendation_type": "no_action",
            "confidence": 0.5
        }

# ==================================================
# FILE: C:\dev\agi-system\src\core\agent\skill_executor.py
# ==================================================

from src.core.agent.skills.create_docx import create_docx


class SkillExecutor:
    """
    Executes approved skills with validated arguments.
    """

    def execute(self, skill_name: str, arguments: dict) -> dict:
        if skill_name == "create_docx":
            return create_docx(**arguments)

        raise ValueError(f"Unknown skill: {skill_name}")

# ==================================================
# FILE: C:\dev\agi-system\src\core\agent\skills\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\agent\skills\create_docx.py
# ==================================================

from docx import Document
from pathlib import Path
from datetime import datetime


def create_docx(title: str, sections: list[str]) -> dict:
    """
    Create a Word document with a title and section headings.
    Returns metadata about the created artifact.
    """

    artifacts_dir = Path("artifacts")
    artifacts_dir.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{title.replace(' ', '_')}_{timestamp}.docx"
    filepath = artifacts_dir / filename

    doc = Document()
    doc.add_heading(title, level=1)

    for section in sections:
        doc.add_heading(section, level=2)
        doc.add_paragraph("")

    doc.save(filepath)

    return {
        "artifact_type": "docx",
        "path": str(filepath),
        "title": title,
        "sections": sections
    }

# ==================================================
# FILE: C:\dev\agi-system\src\core\architecture\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\architecture\agi_system_dataclasses.py
# ==================================================

"""AGI-System core dataclasses (frozen architecture snapshot).

This file implements the *minimum viable* data structures implied by the
architecture Sections 3–11. The intent is to provide stable contracts for
module stubs (NLP -> Executive -> Planner -> Executive -> Executor) and for
cross-cutting subsystems (Persistence, Replay, Error Reporting, Metrics, CMB).

Notes
-----
* All fields include end-of-line comments describing purpose.
* Methods are intentionally minimal (serialization, small helpers).
* Business logic belongs in modules, not dataclasses.
"""

from __future__ import annotations

import time
import uuid
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Optional, List


# -----------------------------
# Common small types / enums
# -----------------------------


class WorkStatus(str, Enum):
    CREATED = "created"
    ACTIVE = "active"
    SUSPENDED = "suspended"
    COMPLETED = "completed"
    FAILED = "failed"
    ABORTED = "aborted"
    EXPIRED = "expired"


class TaskStatus(str, Enum):
    CREATED = "created"
    READY = "ready"
    EXECUTING = "executing"
    SUSPENDED = "suspended"
    COMPLETED = "completed"
    FAILED = "failed"
    ABORTED = "aborted"


class EventType(str, Enum):
    INITIATING = "initiating"
    CONTEXT = "context"
    DECISION = "decision"
    EXECUTION = "execution"
    INTERACTION = "interaction"
    STATE_TRANSITION = "state_transition"
    ERROR = "error"
    OUTCOME = "outcome"
    REFLECTION = "reflection"


class DecisionOutcome(str, Enum):
    ACT = "act"
    DEFER = "defer"
    IGNORE = "ignore"
    ESCALATE = "escalate"
    ABORT = "abort"


class ErrorClass(str, Enum):
    INPUT_VALIDATION = "input_validation"
    TRANSPORT = "transport"
    EXECUTION = "execution"
    EXTERNAL = "external"
    POLICY_SAFETY = "policy_safety"

class DirectiveItemType(Enum):
    STATEMENT = "statement"
    INSTRUCTION = "instruction"
    QUESTION = "question"
    CONSTRAINT = "constraint"
    ASSUMPTION = "assumption"
    CLARIFICATION_NEEDED = "clarification_needed"

class PlanStatus(Enum):
    PROPOSED = "proposed"      # Created by Planner
    APPROVED = "approved"      # Accepted by Executive
    REJECTED = "rejected"      # Planner output not accepted
    SUPERSEDED = "superseded"  # Replaced by a newer plan


def _now() -> float:
    return time.time()


def new_id(prefix: str) -> str:
    """Generate a readable unique identifier."""
    return f"{prefix}_{uuid.uuid4().hex}"


# -----------------------------
# Section 3 — Work
# -----------------------------

@dataclass
class PlanStep:
    step_id: str
    description: str

    # High-level action this step represents
    action_type: str  # e.g. "analyze", "generate", "query", "control"

    # Dependencies (step_ids that must complete first)
    depends_on: List[str] = field(default_factory=list)

    # Optional constraints or hints
    constraints: List[str] = field(default_factory=list)

@dataclass
class Plan:
    """
    Represents a proposed execution plan for a WorkInstance.
    Produced by the Planner and reviewed/approved by the Executive.
    """

    # Identity
    plan_id: str                  # Unique identifier for this plan
    work_id: str                  # WID this plan is associated with

    # Provenance
    created_by: str               # Planner module name or ID
    created_at: float = field(default_factory=time.time)

    # Planning result
    steps: List[PlanStep] = field(default_factory=list)

    # Status managed by Executive
    status: PlanStatus = PlanStatus.PROPOSED

    # Rationale / explanation
    rationale: Optional[str] = None


@dataclass
class WorkInstance:
    """A bounded unit of work (project-like) executed by the system."""

    wid: str = field(default_factory=lambda: new_id("wid"))  # Work instance identifier (primary grouping key)
    title: str = ""  # Human-readable short title for UI/diagnostics
    intent_summary: str = ""  # One-paragraph summary of what the work aims to achieve
    created_at: float = field(default_factory=_now)  # Creation timestamp (commit-to-work moment)
    updated_at: float = field(default_factory=_now)  # Last update timestamp for the work record
    status: WorkStatus = WorkStatus.CREATED  # Current lifecycle state
    priority: int = 50  # Relative urgency (lower or higher semantics chosen by policy; keep consistent system-wide)

    source: str = ""  # Who/what triggered the work (user, subsystem, sensor, scheduler)
    context: dict[str, Any] = field(default_factory=dict)  # Work-scoped context snapshot (inputs, constraints, environment)

    directive_id: Optional[str] = None  # Link back to DirectiveDerivative.did when work originated from a directive
    plan: list["TaskSpec"] = field(default_factory=list)  # Planner-produced task plan (specs)
    tasks: dict[str, "TaskRecord"] = field(default_factory=dict)  # Runtime task records keyed by task_id

    outcome: Optional["OutcomeRecord"] = None  # Terminal evaluation record (success/failure/etc.)

    def touch(self) -> None:
        """Update updated_at to current time."""
        self.updated_at = _now()


# -----------------------------
# NLP -> Executive bridge (Required initial structure)
# -----------------------------

@dataclass
class DirectiveItem:
    item_id: str
    item_type: DirectiveItemType
    text: str

    # Confidence score from NLP / LLM
    confidence: float

    # True if this item blocks work creation
    blocking: bool = False

    # Optional reference to source span in original directive
    source_span: Optional[str] = None


@dataclass
class DirectiveDerivative:
    """Structured interpretation of a human directive produced by NLP/LLM."""

    did: str = field(default_factory=lambda: new_id("did"))  # Directive derivative identifier
    raw_directive: str = ""  # Original text provided by the human (verbatim)
    captured_at: float = field(default_factory=_now)  # Timestamp when directive was received

    user_context: dict[str, Any] = field(default_factory=dict)  # Conversation/user context used during interpretation
    extracted_questions: list[str] = field(default_factory=list)  # Questions contained in the directive (user asking system)
    extracted_statements: list[str] = field(default_factory=list)  # Declarative statements contained in the directive
    extracted_instructions: list[str] = field(default_factory=list)  # Imperatives/commands inferred from the directive

    clarification_questions: list[str] = field(default_factory=list)  # Questions the system needs answered before acting
    assumptions: list[str] = field(default_factory=list)  # Assumptions made if clarification is missing
    constraints: dict[str, Any] = field(default_factory=dict)  # Parsed constraints (time, scope, formatting, safety)

    llm_model: str = ""  # Model identifier used for interpretation (for auditing)
    llm_prompt_version: str = ""  # Prompt/template version used to drive extraction
    llm_raw_response: str = ""  # Optional: raw model text (may be omitted for privacy)
    confidence: float = 0.0  # 0..1 confidence score from NLP pipeline (heuristic or model-provided)

    def needs_clarification(self) -> bool:
        """True if the derivative contains unresolved clarification questions."""
        return len(self.clarification_questions) > 0


# -----------------------------
# Section 4 — Events
# -----------------------------


@dataclass
class EventRecord:
    """An immutable fact recorded in the ordered event stream for a work instance."""

    wid: str  # Work identifier to which this event belongs
    eid: str = field(default_factory=lambda: new_id("eid"))  # Event identifier (unique)
    esn: int = 0  # Event sequence number (monotonic within wid; authoritative ordering)

    event_type: EventType = EventType.CONTEXT  # Category used for analysis and filtering
    timestamp: float = field(default_factory=_now)  # Wall-clock time for performance analysis

    summary: str = ""  # Human-readable one-line summary
    payload: dict[str, Any] = field(default_factory=dict)  # Structured details for machines

    cause_eids: list[str] = field(default_factory=list)  # Explicit causal links to earlier events
    module: str = ""  # Module that emitted/recorded the event
    severity: str = "info"  # Simple severity tag (info|warn|error|critical)


# -----------------------------
# Section 5 — Decisions & Behaviors
# -----------------------------


@dataclass
class DecisionRecord:
    """A recorded moment of choice made within a work instance."""

    wid: str  # Work identifier
    decision_id: str = field(default_factory=lambda: new_id("dec"))  # Decision identifier
    triggering_eids: list[str] = field(default_factory=list)  # Events that prompted this decision

    outcome: DecisionOutcome = DecisionOutcome.ACT  # Decision result (act/defer/ignore/escalate/abort)
    rationale: str = ""  # Human-readable explanation of why the choice was made
    considered_options: list[str] = field(default_factory=list)  # Options considered (names/labels)

    selected_behavior_id: Optional[str] = None  # Behavior chosen (if outcome == ACT)
    created_at: float = field(default_factory=_now)  # Timestamp when decision was made
    module: str = ""  # Decision owner module (typically Executive)


@dataclass
class BehaviorSpec:
    """A reusable behavior definition (metadata + requirements)."""

    behavior_id: str = field(default_factory=lambda: new_id("beh"))  # Behavior identifier
    name: str = ""  # Human-readable behavior name
    description: str = ""  # What the behavior does and when it is applicable

    version: str = "1.0"  # Behavior spec version (semantic versioning recommended)
    required_inputs: list[str] = field(default_factory=list)  # Named inputs required to run this behavior
    produces_outputs: list[str] = field(default_factory=list)  # Named outputs produced by this behavior
    safety_notes: list[str] = field(default_factory=list)  # Safety/policy considerations

    implementation_ref: str = ""  # Pointer to implementation (module:function, plugin id, etc.)


# -----------------------------
# Section 6 — Execution (Tasks)
# -----------------------------


@dataclass
class TaskSpec:
    """Planner-produced plan element describing an intended task."""

    task_id: str = field(default_factory=lambda: new_id("task"))  # Task identifier
    wid: str = ""  # Work identifier that this task belongs to

    behavior_id: str = ""  # Behavior to execute for this task
    name: str = ""  # Human-readable task name
    description: str = ""  # Task intent and expected result

    inputs: dict[str, Any] = field(default_factory=dict)  # Inputs for the behavior/task
    depends_on: list[str] = field(default_factory=list)  # Task IDs that must complete before this one starts
    priority: int = 50  # Task-level priority relative to other tasks in the same work

    ttl_s: Optional[float] = None  # Optional time-to-live; beyond this task is considered expired


@dataclass
class TaskRecord:
    """Runtime record of a task execution."""

    task_id: str  # Task identifier
    wid: str  # Work identifier

    status: TaskStatus = TaskStatus.CREATED  # Current runtime state
    created_at: float = field(default_factory=_now)  # When the task record was created
    started_at: Optional[float] = None  # When execution began
    ended_at: Optional[float] = None  # When execution completed/failed/aborted

    executor_module: str = ""  # Which module/executor is running the task
    attempt: int = 0  # Execution attempt counter (increments on retry)

    last_error_id: Optional[str] = None  # Link to ErrorRecord.error_id for last failure
    outputs: dict[str, Any] = field(default_factory=dict)  # Outputs produced by the task

    def mark_started(self) -> None:
        self.status = TaskStatus.EXECUTING
        self.started_at = _now()

    def mark_completed(self, outputs: Optional[dict[str, Any]] = None) -> None:
        self.status = TaskStatus.COMPLETED
        self.ended_at = _now()
        if outputs:
            self.outputs.update(outputs)

    def mark_failed(self, error_id: str) -> None:
        self.status = TaskStatus.FAILED
        self.ended_at = _now()
        self.last_error_id = error_id


# -----------------------------
# Outcome (Sections 3/4/10)
# -----------------------------


@dataclass
class OutcomeRecord:
    """Terminal evaluation for a work instance."""

    wid: str  # Work identifier
    outcome_id: str = field(default_factory=lambda: new_id("out"))  # Outcome record identifier

    status: WorkStatus = WorkStatus.COMPLETED  # Terminal status
    summary: str = ""  # Human-readable explanation of the outcome
    metrics: dict[str, Any] = field(default_factory=dict)  # Outcome metrics (latency, quality scores, etc.)

    completed_at: float = field(default_factory=_now)  # Timestamp when outcome was recorded


# -----------------------------
# Section 7 — Identity & Traceability (Transaction)
# -----------------------------


@dataclass
class TransportHeader:
    """Transport-level metadata used by CMB infrastructure."""

    xid: str = field(default_factory=lambda: new_id("xid"))  # Transport transaction identifier
    channel: str = ""  # Logical CMB channel name
    require_ack: bool = True  # Whether infrastructure must produce/await acknowledgment
    ttl_s: float = 10.0  # Delivery time-to-live for retries/expiration


@dataclass
class SemanticHeader:
    """Semantic metadata used by modules to correlate messages to work/event/task context."""

    wid: Optional[str] = None  # Work identifier
    eid: Optional[str] = None  # Event identifier associated with this message
    esn: Optional[int] = None  # Event sequence number (if message represents an event)

    task_id: Optional[str] = None  # Task identifier
    decision_id: Optional[str] = None  # Decision identifier

    correlation_id: Optional[str] = None  # Cross-message correlation key (request/response pairing)

    source: str = ""  # Module sending the message
    targets: list[str] = field(default_factory=list)  # Intended recipients
    message_type: str = ""  # Semantic type label (command/event/ack/error/etc.)


@dataclass
class MessageEnvelope:
    """The unit of communication on the CMB (transport + semantic + payload)."""

    message_id: str = field(default_factory=lambda: new_id("msg"))  # Unique message identifier
    transport: TransportHeader = field(default_factory=TransportHeader)  # Transport metadata
    semantic: SemanticHeader = field(default_factory=SemanticHeader)  # Semantic metadata
    payload: dict[str, Any] = field(default_factory=dict)  # Message payload (content)
    created_at: float = field(default_factory=_now)  # Timestamp when message was created

    def to_dict(self) -> dict[str, Any]:
        """Shallow serialization helper for logging/persistence."""
        return {
            "message_id": self.message_id,
            "created_at": self.created_at,
            "transport": self.transport.__dict__,
            "semantic": self.semantic.__dict__,
            "payload": self.payload,
        }


@dataclass
class TransactionRecord:
    """Tracks transport-level lifecycle for a message exchange (ACK/timeout/retry)."""

    xid: str  # Transaction identifier (matches TransportHeader.xid)
    message_id: str  # Message identifier being tracked
    channel: str  # Channel used

    created_at: float = field(default_factory=_now)  # When tracking began
    last_attempt_at: Optional[float] = None  # Time of last send attempt
    attempts: int = 0  # Number of send attempts

    acked: bool = False  # Whether an ACK was received
    acked_at: Optional[float] = None  # When the ACK was received
    expired: bool = False  # Whether TTL was exceeded
    error: Optional[str] = None  # Transport error detail, if any


# -----------------------------
# Section 10 — Errors
# -----------------------------


@dataclass
class ErrorRecord:
    """Structured error record for system-wide reporting and persistence."""

    error_id: str = field(default_factory=lambda: new_id("err"))  # Unique error identifier
    error_class: ErrorClass = ErrorClass.EXECUTION  # High-level classification

    wid: Optional[str] = None  # Work identifier impacted by this error
    task_id: Optional[str] = None  # Task identifier impacted by this error
    xid: Optional[str] = None  # Transaction identifier (if transport-related)

    module: str = ""  # Module where the error was detected
    message: str = ""  # Human-readable error message
    details: dict[str, Any] = field(default_factory=dict)  # Structured diagnostic details

    created_at: float = field(default_factory=_now)  # Time error was recorded
    severity: str = "error"  # Severity tag (warn|error|critical)


# -----------------------------
# Section 9 — Replay
# -----------------------------


@dataclass
class ReplayRequest:
    """Request to replay a work instance or segment in a safe environment."""

    replay_id: str = field(default_factory=lambda: new_id("replay"))  # Replay request identifier
    wid: str = ""  # Work identifier to replay

    mode: str = "full"  # Replay mode: full|partial|simulate
    from_esn: Optional[int] = None  # Starting sequence number (partial replay)
    to_esn: Optional[int] = None  # Ending sequence number (partial replay)

    simulate_side_effects: bool = False  # If False, external side effects should be disabled/mocked
    overrides: dict[str, Any] = field(default_factory=dict)  # What-if overrides (decision policies, behavior selection)

    created_at: float = field(default_factory=_now)  # When replay was requested
    requested_by: str = ""  # Module/user requesting the replay


# -----------------------------
# Section 11 — Learning / Behavior extraction
# -----------------------------


@dataclass
class BehaviorCandidate:
    """Candidate behavior extracted from historical work instances."""

    candidate_id: str = field(default_factory=lambda: new_id("cand"))  # Candidate identifier
    name: str = ""  # Candidate name
    description: str = ""  # Why it is useful / what it does

    source_wids: list[str] = field(default_factory=list)  # Work instances supporting this pattern
    applicability_conditions: dict[str, Any] = field(default_factory=dict)  # When it should be applied
    expected_outputs: list[str] = field(default_factory=list)  # Outputs expected when it succeeds
    known_failure_modes: list[str] = field(default_factory=list)  # Observed failure patterns

    metrics: dict[str, Any] = field(default_factory=dict)  # Evidence metrics (success rate, latency)
    status: str = "candidate"  # candidate|validated|approved|rejected

    created_at: float = field(default_factory=_now)  # When candidate was created


@dataclass
class BehaviorLibraryEntry:
    """Approved behavior registered in the behavior library."""

    behavior: BehaviorSpec  # Behavior specification
    approved_at: float = field(default_factory=_now)  # When it was approved
    approved_by: str = ""  # Who/what approved it (human, policy engine)
    provenance_candidate_id: Optional[str] = None  # Link to BehaviorCandidate that produced it


@dataclass
class MetricSample:
    """A single metric observation tied to work/task/behavior."""

    metric_id: str = field(default_factory=lambda: new_id("met"))  # Metric identifier
    name: str = ""  # Metric name (latency_ms, retry_count, etc.)
    value: float = 0.0  # Numeric value
    unit: str = ""  # Unit label (ms, count, score)

    wid: Optional[str] = None  # Work identifier
    task_id: Optional[str] = None  # Task identifier
    behavior_id: Optional[str] = None  # Behavior identifier
    timestamp: float = field(default_factory=_now)  # Time captured

# ==================================================
# FILE: C:\dev\agi-system\src\core\behaviors\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\behaviors\behavior_stub.py
# ==================================================

"""
Module: behavior_stub.py
Location: src/core/behaviors/
Version: 0.1.0

Simulates the Behavior Module in the AGI system. Subscribes to the Control Channel (CC)
and logs received messages. This stub allows for testing message routing from the executive.

Depends on: module_endpoint >= 0.1.0, cognitive_message >= 0.1.0, cmb_channel_config >= 0.1.0
"""

import zmq
import time
from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.cmb.channel_registry import ChannelRegistry
from src.core.cmb.endpoint_config import MultiChannelEndpointConfig
from src.core.messages.ack_message import AckMessage


def main():

    endpoint = None
    module_id = "behavior"
    logger = print  # Simple logger function

    # Build ChannelRegistry once
    #ChannelRegistry.initialize()

    # Decide which channels the GUI participates in.
    # Start minimal for the demo: choose the channel(s) you use in the dropdown.
    gui_channels = ["CC", "SMC", "VB", "BFC", "DAC", "EIG", "PC", "MC", "IC", "TC"]

    cfg = MultiChannelEndpointConfig.from_channel_names(
        module_id=module_id,  # Identity of this module on the bus
        channel_names=gui_channels,
        host="localhost",
        poll_timeout_ms=50,
    )

    # Create endpoint (logger uses your GUI log function)
    endpoint = ModuleEndpoint(
        config=cfg,
        logger=print,
        serializer=lambda x: x if isinstance(x, (bytes, bytearray)) else str(x).encode("utf-8"),
        deserializer=lambda b: b,  # Keep bytes; GUI will parse ACK vs MSG
    )

    endpoint.start()
    logger("[BehaviorStub] Endpoint started.")

    try:
        print("[BehaviorStub] Listening for control messages on queue..")
        while True:
                    
            raw_msg = endpoint.recv(timeout=5.0)

            if raw_msg is None:
                continue

            msg = raw_msg

            logger(f"[BehaviorStub] Received {msg.msg_type} from {msg.source}")
            logger(f"[BehaviorStub] Payload: {msg.payload}")

            
            # send ACK back
            try:
                ack = AckMessage.create(
                    msg_type="ACK",
                    ack_type="MESSAGE_DELIVERED_ACK",
                    status="SUCCESS",
                    source="behavior",
                    targets=[msg.source],
                    correlation_id=msg.message_id,
                    payload={ 
                        "status": "published",
                        "message_id": msg.message_id
                    }
                )

                #endpoint.send("CC", msg.source, AckMessage.to_bytes(ack))

                print(
                  f"[BEHAVIOR] Sent ACK to {msg.source} "
                  f"for {msg.message_id}"
                )
            except Exception as e:
                logger(f"[BehaviorStub] ERROR sending ACK: {e}")

            ack_msg = endpoint.recv_ack(timeout=2.0)
            if ack_msg:
                logger(f"[BehaviorStub] Received ACK: {ack_msg.ack_type} for {ack_msg.correlation_id}")

    except KeyboardInterrupt:
        print("[BehaviorStub] Interrupted by user.")

    finally:
        endpoint.stop()
        logger("[BehaviorStub] Shutdown.")


if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\channel_registry.py
# ==================================================

"""
Module: channel_registry.py
Location: src/core/cmb/
Version: 0.2.0

Defines the authoritative registry of Cognitive Message Bus (CMB) channels.
Each channel is described declaratively via ChannelConfig objects.
"""

from dataclasses import dataclass
from enum import Enum
from typing import Dict

import zmq

from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_entry import LogEntry
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink

# ----------------------------
# Inbound delivery semantics
# ----------------------------

class InboundDelivery(Enum):
    """
    Defines how messages are delivered to modules on this channel.
    """
    DIRECTED = "DIRECTED"    # ROUTER -> DEALER (addressed, ACK-capable)
    BROADCAST = "BROADCAST"  # PUB -> SUB (fan-out, no ACKs)
    NONE = "NONE"            # Send-only channel


# ----------------------------
# Channel configuration
# ----------------------------

@dataclass(frozen=True)
class ChannelConfig:
    """
    Declarative configuration for a single CMB channel.
    """

    name: str

    # Outbound path (module -> router)
    router_port: int
    outbound_socket_type: int = zmq.DEALER

    # Inbound semantics
    inbound_delivery: InboundDelivery = InboundDelivery.DIRECTED
    inbound_port: int | None = None

    # ACK path (only meaningful for DIRECTED channels)
    ack_port: int | None = None
    ack_socket_type: int = zmq.DEALER

# ----------------------------
# Legacy port assignments
# ----------------------------

# Dictionary mapping each channel acronym to a TCP port base
# These ports must match bindings in the CMB routers and the endpoints
CMB_CHANNEL_INGRESS_PORTS = {
    "CC":   6001,   # Control Channel
    "SMC":  6002,   # Symbolic Message Channel
    "VB":   6003,   # Vector Bus
    "BFC":  6004,   # Behavioral Flow Channel
    "DAC":  6005,   # Diagnostic and Awareness Channel
    "EIG":  6006,   # External Interface Gateway
    "PC":   6007,   # Perception Channel
    "MC":   6008,   # Memory Channel
    "IC":   6009,   # Introspection Channel
    "TC":   6010    # Threat Channel
}
CMB_CHANNEL_EGRESS_PORTS = {
    "CC":   7001,   # Control Channel
    "SMC":  7002,   # Symbolic Message Channel
    "VB":   7003,   # Vector Bus
    "BFC":  7004,   # Behavioral Flow Channel
    "DAC":  7005,   # Diagnostic and Awareness Channel
    "EIG":  7006,   # External Interface Gateway
    "PC":   7007,   # Perception Channel
    "MC":   7008,   # Memory Channel
    "IC":   7009,   # Introspection Channel
    "TC":   7010    # Threat Channel
}

CMB_ACK_INGRESS_PORTS = {
    "CC":  6101,   # Control Channel
    "SMC": 6101,   # Symbolic Message Channel
    "VB":  6101,   # Vector Bus
    "BFC": 6101,   # Behavioral Flow Channel
    "DAC": 6101,   # Diagnostic and Awareness Channel
    "EIG": 6101,   # External Interface Gateway
    "PC":  6101,   # Perception Channel
    "MC":  6101,   # Memory Channel
    "IC":  6101,   # Introspection Channel
    "TC":  6101    # Threat Channel
}

CMB_ACK_EGRESS_PORTS = {
    "CC":  6102,   # Control Channel
    "SMC": 6102,   # Symbolic Message Channel
    "VB":  6102,   # Vector Bus
    "BFC": 6102,   # Behavioral Flow Channel
    "DAC": 6102,   # Diagnostic and Awareness Channel
    "EIG": 6102,   # External Interface Gateway
    "PC":  6102,   # Perception Channel
    "MC":  6102,   # Memory Channel
    "IC":  6102,   # Introspection Channel
    "TC":  6102    # Threat Channel
}

CMB_ACK_PORT = 6102        # Shared ACK ingress/egress (current policy)
SUBSCRIPTION_OFFSET = 1000


# ----------------------------
# Channel Registry
# ----------------------------

class InboundDelivery(Enum):
    """How inbound messages are delivered to modules."""

    DIRECTED = "directed"      # ROUTER/DEALER identity-addressed
    BROADCAST = "broadcast"    # PUB/SUB fanout (not used in this demo)


class ChannelRegistry:
    """
    Central registry of all CMB channels.
    """

    _channels: Dict[str, ChannelConfig] = {}

    @classmethod
    def initialize(cls) -> None:
        """
        Build the registry.
        Call once at system startup.
        """

        cls._channels = {
            # Control-plane channels (DIRECTED)
            "CC": ChannelConfig(
                name="CC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["CC"],
                inbound_delivery=InboundDelivery.DIRECTED,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["CC"],
                ack_port=CMB_ACK_EGRESS_PORTS["CC"],
            ),

            "SMC": ChannelConfig(
                name="SMC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["SMC"],
                inbound_delivery=InboundDelivery.DIRECTED,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["SMC"],
                ack_port=CMB_ACK_EGRESS_PORTS["SMC"],
            ),

            "VB": ChannelConfig(
                name="VB",
                router_port=CMB_CHANNEL_INGRESS_PORTS["VB"],
                inbound_delivery=InboundDelivery.DIRECTED,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["VB"],
                ack_port=CMB_ACK_PORT,
            ),

            "BFC": ChannelConfig(
                name="BFC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["BFC"],
                inbound_delivery=InboundDelivery.DIRECTED,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["BFC"],
                ack_port=CMB_ACK_PORT,
            ),

            "DAC": ChannelConfig(
                name="DAC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["DAC"],
                inbound_delivery=InboundDelivery.DIRECTED,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["DAC"],
                ack_port=CMB_ACK_PORT,
            ),

            "IC": ChannelConfig(
                name="IC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["IC"],
                inbound_delivery=InboundDelivery.DIRECTED,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["IC"],
                ack_port=CMB_ACK_PORT,
            ),

            "TC": ChannelConfig(
                name="TC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["TC"],
                inbound_delivery=InboundDelivery.DIRECTED,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["TC"],
                ack_port=CMB_ACK_PORT,
            ),

            # Broadcast-style channels
            "PC": ChannelConfig(
                name="PC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["PC"],
                inbound_delivery=InboundDelivery.BROADCAST,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["PC"] + SUBSCRIPTION_OFFSET,
                ack_port=None,
            ),

            "MC": ChannelConfig(
                name="MC",
                router_port=CMB_CHANNEL_INGRESS_PORTS["MC"],
                inbound_delivery=InboundDelivery.BROADCAST,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["MC"] + SUBSCRIPTION_OFFSET,
                ack_port=None,
            ),

            "EIG": ChannelConfig(
                name="EIG",
                router_port=CMB_CHANNEL_INGRESS_PORTS["EIG"],
                inbound_delivery=InboundDelivery.BROADCAST,
                inbound_port=CMB_CHANNEL_EGRESS_PORTS["EIG"] + SUBSCRIPTION_OFFSET,
                ack_port=None,
            ),
        }

    @classmethod
    def get(cls, channel_name: str) -> ChannelConfig:
        
        if channel_name not in cls._channels:
            raise KeyError(f"Unknown CMB channel: {channel_name}")
        return cls._channels[channel_name]

    @classmethod
    def all(cls) -> Dict[str, ChannelConfig]:
        return dict(cls._channels)

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_channel_config.py
# ==================================================

"""
Module: cmb_channel_config.py
Location: src/core/cmb/
Version: 0.1.0

Defines a consistent mapping of Cognitive Message Bus (CMB) channels to their associated TCP ports.
This file should be treated as a central reference for routing and interface connections.
"""

# Dictionary mapping each channel acronym to a TCP port base
# These ports must match bindings in the CMB routers and the endpoints
CMB_CHANNEL_INGRESS_PORTS = {
    "CC":   6001,   # Control Channel
    "SMC":  6002,   # Symbolic Message Channel
    "VB":   6003,   # Vector Bus
    "BFC":  6004,   # Behavioral Flow Channel
    "DAC":  6005,   # Diagnostic and Awareness Channel
    "EIG":  6006,   # External Interface Gateway
    "PC":   6007,   # Perception Channel
    "MC":   6008,   # Memory Channel
    "IC":   6009,   # Introspection Channel
    "TC":   6010    # Threat Channel
}
CMB_CHANNEL_EGRESS_PORTS = {
    "CC":   7001,   # Control Channel
    "SMC":  7002,   # Symbolic Message Channel
    "VB":   7003,   # Vector Bus
    "BFC":  7004,   # Behavioral Flow Channel
    "DAC":  7005,   # Diagnostic and Awareness Channel
    "EIG":  7006,   # External Interface Gateway
    "PC":   7007,   # Perception Channel
    "MC":   7008,   # Memory Channel
    "IC":   7009,   # Introspection Channel
    "TC":   7010    # Threat Channel
}

CMB_ACK_INGRESS_PORTS = {
    "CC":  6101,   # Control Channel
    "SMC": 6101,   # Symbolic Message Channel
    "VB":  6101,   # Vector Bus
    "BFC": 6101,   # Behavioral Flow Channel
    "DAC": 6101,   # Diagnostic and Awareness Channel
    "EIG": 6101,   # External Interface Gateway
    "PC":  6101,   # Perception Channel
    "MC":  6101,   # Memory Channel
    "IC":  6101,   # Introspection Channel
    "TC":  6101    # Threat Channel
}

CMB_ACK_EGRESS_PORTS = {
    "CC":  6102,   # Control Channel
    "SMC": 6102,   # Symbolic Message Channel
    "VB":  6102,   # Vector Bus
    "BFC": 6102,   # Behavioral Flow Channel
    "DAC": 6102,   # Diagnostic and Awareness Channel
    "EIG": 6102,   # External Interface Gateway
    "PC":  6102,   # Perception Channel
    "MC":  6102,   # Memory Channel
    "IC":  6102,   # Introspection Channel
    "TC":  6102    # Threat Channel
}


# Ports offset by 1000 for Subscription channels
def get_subscription_offset():
    return 1000

# Utility function to fetch port by channel
def get_channel_port(channel_name: str) -> int:
    if channel_name not in CMB_CHANNEL_INGRESS_PORTS:
        raise ValueError(f"Unknown CMB channel: {channel_name}")
    return CMB_CHANNEL_INGRESS_PORTS[channel_name]

def get_channel_ingress_port(channel_name: str) -> int:
    if channel_name not in CMB_CHANNEL_INGRESS_PORTS:
        raise ValueError(f"Unknown CMB channel: {channel_name}")
    return CMB_CHANNEL_INGRESS_PORTS[channel_name]

def get_channel_egress_port(channel_name: str) -> int:
    if channel_name not in CMB_CHANNEL_EGRESS_PORTS:
        raise ValueError(f"Unknown CMB channel: {channel_name}")
    return CMB_CHANNEL_EGRESS_PORTS[channel_name]

def get_ack_ingress_port(channel_name: str) -> int:
    if channel_name not in CMB_ACK_INGRESS_PORTS:
        raise ValueError(f"Unknown CMB channel: {channel_name}")
    return CMB_ACK_INGRESS_PORTS[channel_name] 

def get_ack_port(channel_name: str) -> int:
    if channel_name not in CMB_ACK_EGRESS_PORTS:
        raise ValueError(f"Unknown CMB channel: {channel_name}")
    return CMB_ACK_EGRESS_PORTS[channel_name] 

def get_ack_egress_port(channel_name: str) -> int:
    if channel_name not in CMB_ACK_EGRESS_PORTS:
        raise ValueError(f"Unknown CMB channel: {channel_name}")
    return CMB_ACK_EGRESS_PORTS[channel_name] 

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_channel_router_port.py
# ==================================================

"""
Location: src/core/cmb/
Version: 0.1.0

Defines the ChannelRouterPort class, which abstracts the interface between a module and the router for a specific CMB channel.
Handles sending messages via PUSH sockets and receiving via SUB sockets for specified topics.
Depends on: cognitive_message >= 0.1.0
"""

from multiprocessing import context
import zmq
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.cmb.cmb_channel_config import get_channel_port
class ChannelRouterPort:
    def __init__(self, module_name: str, channel_name: str):

        self.module_name = module_name
        self.channel_name = channel_name
        self.port_number = get_channel_port(channel_name)
  

        # Create Dealer socket for sending messages to router
        self.context = zmq.Context()
        self.router_socket = self.context.socket(zmq.DEALER)
        self.router_socket.setsockopt_string(zmq.IDENTITY, module_name)
        self.router_socket.connect(f"tcp://localhost:{self.port_number}")

    def send(self, message: CognitiveMessage):
        assert isinstance(message, CognitiveMessage)
        self.router_socket.send_multipart([
        message.source.encode(),
        message.to_bytes()
    ])
        print(f"Sent message from {message.source} to {message.targets}")


    def receive(self) -> CognitiveMessage:
        topic, raw_msg = self.router_socket.recv_multipart()
        return CognitiveMessage.from_bytes(raw_msg)

    def close(self):
        self.router_socket.close()
        self.context.term()

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_exceptions.py
# ==================================================

class TransportError(Exception):
    def __init__(self, tx_id, reason, details=None):
        self.tx_id = tx_id
        self.reason = reason
        self.details = details
        super().__init__(reason)

class AckTimeoutError(TransportError):
    pass

class InvalidAckError(TransportError):
    pass

class ProtocolViolationError(TransportError):
    pass

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_router.py
# ==================================================

"""cmb_router.py

Version: demo-fork

ROUTER-based channel router for the Cognitive Message Bus (CMB).

- One router per channel
- Ingress ROUTER: receives from module outbound DEALER sockets
- Module egress ROUTER: forwards to module inbound DEALER sockets
- ACK egress ROUTER: forwards ACKs to module ACK DEALER sockets
- Sends immediate ROUTER_ACK **only for non-ACK messages**

This is a lightly corrected version of your current router to avoid emitting
ROUTER_ACK for ACK messages (which can create ack-of-ack loops) and to avoid
referencing an undefined `msg` when processing ACKs.
"""

from __future__ import annotations

import json
import threading
from core.cmb.channel_registry import ChannelRegistry
import zmq

from src.core.messages.cognitive_message import CognitiveMessage
from src.core.messages.ack_message import AckMessage
from src.core.cmb.cmb_channel_config import (
    get_channel_ingress_port,
    get_ack_egress_port,
    get_channel_egress_port,
)

from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_entry import LogEntry
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink

class ChannelRouter:
    def __init__(self, channel_name: str, host: str = "localhost"):
        self.channel_name = channel_name
        self.host = host

        self.router_port = get_channel_ingress_port(channel_name)
        self.module_egress_port = get_channel_egress_port(channel_name)
        self.ack_port = get_ack_egress_port(channel_name)

        self._stop_evt = threading.Event()
        self._thread = None

        # Logging
        self.log_manager = LogManager(min_severity=LogSeverity.INFO)
        self.log_manager.register_sink(
        FileLogSink("logs/system.jsonl")
        )

        self.logger = Logger(self.channel_name, self.log_manager)

        self.logger.info(
            event_type="ROUTER_INIT",
            message=f"Router {self.channel_name} port: {self.router_port} ack: {self.ack_port} module_egress: {self.module_egress_port}",
            payload={
                "note": "no payload"
            }
        )
        
    def start(self) -> None:
        if self._thread and self._thread.is_alive():
            return
        self._stop_evt.clear()
        self._thread = threading.Thread(
            target=self._run,
            name=f"ChannelRouter[{self.channel_name}]",
            daemon=False,
        )
        self._thread.start()
        
        self.logger.info(
                    event_type="ROUTER_START",
                    message=f"Router {self.channel_name} started",
                    payload={
                        "note": "no payload"
                    }
                )

    def stop(self) -> None:
        self._stop_evt.set()
        if self._thread:
            self._thread.join(timeout=2.0)
        
        self.logger.info(
                event_type="ROUTER_STOP",
                message=f"Router {self.channel_name} stopped",
                payload={
                    "note": "no payload"
                }
            )

    def _run(self) -> None:
        ctx = zmq.Context.instance()
        router_sock = ctx.socket(zmq.ROUTER)
        router_sock.bind(f"tcp://{self.host}:{self.router_port}")

        module_egress_sock = ctx.socket(zmq.ROUTER)
        module_egress_sock.bind(f"tcp://{self.host}:{self.module_egress_port}")

        ack_sock = ctx.socket(zmq.ROUTER)
        ack_sock.bind(f"tcp://{self.host}:{self.ack_port}")

        poller = zmq.Poller()
        poller.register(router_sock, zmq.POLLIN)

        self.logger.info(
                event_type="ROUTER_START_RUN",
                message=f"[Router.{self.channel_name}] ROUTER ingress on {self.router_port}, egress on {self.module_egress_port}, ACK on {self.ack_port}",
                payload={
                    "note": "no payload"
                }
            )

        try:
            while not self._stop_evt.is_set():
                events = dict(poller.poll(100))
                if router_sock not in events:
                    continue

                frames = router_sock.recv_multipart()
                sender_id = frames[0]
                payload = frames[-1]

                try:
                    obj = json.loads(payload.decode("utf-8"))
                except Exception as e:
                
                    self.logger.info(
                        event_type="ROUTER_EXCEPTIOM_ERROR",
                        message=f"Invalid JSON message: {e}",
                        payload={
                            "note": "no payload"
                        }
                    )

                    continue

                msg_type = obj.get("msg_type")

                # --- ACK messages: forward only ---
                if msg_type == "ACK":
                    try:
                        ack = AckMessage.from_dict(obj)
                    except Exception as e:
                        
                        self.logger.info(
                            event_type="ROUTER_EXCEPTIOM_ERROR",
                            message=f"[Router.{self.channel_name} ERROR] Invalid ACK message: {e}",
                            payload={
                                "note": "no payload"
                            }
                        )

                        continue

                    if not ack.targets:
                       
                        self.logger.info(
                            event_type="ROUTER_NO_ACK_TARGETS_ERROR",
                            message=f"[Router.{self.channel_name} ERROR] ACK has no targets",
                            payload={
                                "note": "no payload"
                            }
                        )

                        continue

                    dest = ack.targets[0].encode("utf-8")
                    ack_sock.send_multipart([dest, b"", payload])
                    continue

                # --- Non-ACK messages: forward to targets + emit ROUTER_ACK ---
                try:
                    msg = CognitiveMessage.from_dict(obj)
                except Exception as e:
                    
                    self.logger.info(
                            event_type="ROUTER_INVALID_MESSAGE_ERROR",
                            message=f"[Router.{self.channel_name} invalid message not Cognitive Message {e}",
                            payload={
                                "note": "no payload"
                            }
                        )
                    
                    continue

                for target in msg.targets:
                    module_egress_sock.send_multipart([
                        target.encode("utf-8"),
                        b"",
                        payload,
                    ])

                # Immediate ROUTER_ACK to the sender (logical sender = msg.source)
                router_ack = AckMessage.create(
                    msg_type="ACK",
                    ack_type="ROUTER_ACK",
                    status="SUCCESS",
                    source="CMB_ROUTER",
                    targets=[msg.source],
                    correlation_id=msg.message_id,
                    payload={
                        "channel": self.channel_name,
                        "status": "published",
                        "message_id": msg.message_id,
                    },
                )

                ack_sock.send_multipart([
                    sender_id,
                    b"",
                    router_ack.to_bytes(),
                ])

        finally:
            router_sock.close()
            module_egress_sock.close()
            ack_sock.close()
            # Do not ctx.term() when using Context.instance() in multi-thread/process environments
            
            self.logger.info(
                            event_type="ROUTER_SHUTDOWN_COMPLETE",
                            message=f"[Router.{self.channel_name}] shutdown complete",
                            payload={
                                "note": "no payload"
                            }
                        )

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_router_entry.py
# ==================================================

# src/core/cmb/cmb_router_entry.py

import argparse
from src.core.cmb.cmb_router import ChannelRouter  # or whatever your class is named

def main():
    parser = argparse.ArgumentParser(description="CMB Channel Router")
    parser.add_argument("--channel", required=True, help="Channel name (e.g. CC, VB)")
    args = parser.parse_args()

    router = ChannelRouter(channel_name=args.channel)
    router.start()   # or start(), loop(), etc.
    

if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_router_legacy_3.py
# ==================================================

"""
Module: cmb_router.py
Location: src/core/cmb/
Version: 0.2.0

Minimal ROUTER-based channel router for the Cognitive Message Bus.

- One router per channel
- ROUTER → ROUTER forwarding
- Immediate ROUTER_ACK
- No state machine
"""

import json
from core.messages.message_module import MessageType
import zmq
import threading
import time
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.messages.ack_message import AckMessage
from src.core.cmb.cmb_channel_config import (
    get_channel_ingress_port,
    get_ack_egress_port,
    get_channel_egress_port
)


class ChannelRouter:
    def __init__(self, channel_name: str, host: str = "localhost"):
        self.channel_name = channel_name
        self.host = host

        self.router_port = get_channel_ingress_port(channel_name)
        self.module_egress_port = get_channel_egress_port(channel_name)
        self.ack_port = get_ack_egress_port(channel_name)

        self._stop_evt = threading.Event()
        self._thread = None
        print(f"[ROUTER.{self.channel_name}] router_port = {self.router_port}, ack_port = {self.ack_port}, module_egress_port = {self.module_egress_port}")

    def start(self) -> None:
        if self._thread and self._thread.is_alive():
            return
        self._stop_evt.clear()
        self._thread = threading.Thread(
            target=self._run,
            name=f"ChannelRouter[{self.channel_name}]",
            daemon=False,
        )
        self._thread.start()
        print(f"[Router.{self.channel_name}] Channel '{self.channel_name}' started")

    def stop(self) -> None:
        self._stop_evt.set()
        if self._thread:
            self._thread.join(timeout=2.0)
        print(f"[Router.{self.channel_name}] Channel '{self.channel_name}' stopped")


    def _run(self) -> None:
        ctx = zmq.Context.instance()
        print(f"[Router.{self.channel_name}] starting on {self.router_port}")
        router_sock = ctx.socket(zmq.ROUTER)
        router_sock.bind(f"tcp://{self.host}:{self.router_port}")

        module_egress_sock = ctx.socket(zmq.ROUTER)
        module_egress_sock.bind(f"tcp://{self.host}:{self.module_egress_port}")

        ack_sock = ctx.socket(zmq.ROUTER)
        ack_sock.bind(f"tcp://{self.host}:{self.ack_port}")

        poller = zmq.Poller()
        poller.register(router_sock, zmq.POLLIN)
        poller.register(module_egress_sock, zmq.POLLIN)

        print(f"[Router.{self.channel_name}] {self.channel_name} ROUTER on {self.router_port}, ")
        print(f"[Router.{self.channel_name}] {self.channel_name} ACK on {self.ack_port}")

        try:
            while not self._stop_evt.is_set():
                events = dict(poller.poll(100))

                if router_sock not in events:
                    continue
                
                # Receive message
                frames = router_sock.recv_multipart()
                sender_id = frames[0]
                payload = frames[-1]

                # ---- Parse message ----
                try:
                    obj = json.loads(payload.decode("utf-8"))
                except Exception as e:    
                    print(f"[Router.{self.channel_name} ERROR] Invalid JSON message: {e}")
                    continue

                # ---- Determine message type ----
                msg_type = obj.get("msg_type")
                print(f"[Router.{self.channel_name} DEBUG] Received msg_type={msg_type!r} from sender_id={sender_id!r}")

                # ---- Process COMMAND messages ----
                if msg_type != "ACK":
                    try:
                        msg = CognitiveMessage.from_dict(obj)
                    except Exception as e:
                        print(f"[Router.{self.channel_name} ERROR] Invalid COMMAND message: {e}")
                        continue

                    # ---- Forward to targets ----
                    for target in msg.targets:
                        print(f"[Router.{self.channel_name} DEBUG] Sending {msg.msg_type} to identity={target!r}")

                        module_egress_sock.send_multipart([
                            target.encode("utf-8"),
                            b"",
                            payload,
                        ])
                else:
                    # ---- Process ACK messages ----
                    # This handles the ACK forwarding for MESSAGE_DELIVERED_ACK
                    ack = AckMessage.from_dict(obj)
                    print(f"[Router.{self.channel_name} DEBUG] Forwarding {ack.ack_type} id: {ack.message_id} to {ack.targets[0]!r}")

                    # ---- Forward ACK to destination ----
                    dest = ack.targets[0].encode("utf-8")
                    ack_sock.send_multipart([dest, b"", payload])        

                # ---- Immediate ROUTER_ACK ----
                ack = AckMessage.create(
                    msg_type="ACK",
                    ack_type="ROUTER_ACK",
                    status="SUCCESS",
                    source="CMB_ROUTER",
                    targets=[msg.source],
                    correlation_id=msg.message_id,
                    payload={
                        "channel": self.channel_name,
                        "status": "published",
                        "message_id": msg.message_id,
                    },
                )

                ack_sock.send_multipart([
                    sender_id,
                    b"",
                    AckMessage.to_bytes(ack),
                ])

                print(f"[Router.{self.channel_name}] sent ROUTER_ACK to {sender_id.decode('utf-8')} for {msg.message_id}")

        finally:
            router_sock.close()
            ack_sock.close()
            ctx.term()
            print(f"[Router.{self.channel_name}] {self.channel_name} shutdown complete")

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_router_legacy_pubsub.py
# ==================================================

"""
Module: cmb_router.py
Location: src/core/cmb/
Version: 0.1.0

Implements a ZMQ-based message router for a single CMB channel.
Receives messages via ROUTER or PULL socket and republishes to subscribers via PUB socket.
This router is channel-agnostic and can be launched per CMB channel.

Depends: cognitive_message >= 0.1.0, module_endpoint >= 0.1.0, cmb_channel_config >= 0.1.0
"""

import zmq
import threading
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.messages.ack_message import AckMessage
from src.core.cmb.cmb_channel_config import get_channel_port, get_channel_publish_port
from src.core.cmb.cmb_channel_config import get_ack_Ingress_port, get_ack_Egress_port
import uuid 


class CMBRouter:
    def __init__(self, channel_name: str):
        self.channel_name = channel_name
        self.port_in  = get_channel_port(channel_name) # ROUTER/PULL input port
        self.port_out = get_channel_publish_port(channel_name)
        self.port_ack_ingress = get_ack_Ingress_port(channel_name)    # ROUTER  port for ACKs
        self.port_ack_egress  = get_ack_Egress_port(channel_name) # PUB port for ACKs

    
    def validate_message(self, msg: CognitiveMessage) -> bool:
          if not msg.source or not msg.targets:
             return False
          return True
    @staticmethod
    def system_ack(payload: dict) -> "CognitiveMessage":
        return CognitiveMessage(
            message_id=str(uuid.uuid4()),
            source="CMB_ROUTER",
            tartegs=[],
            payload=payload,
            priority=0
        )
    
    def start(self):
        print(f"[CMBRouter] Starting router for channel '{self.channel_name}' on port 6000")
        thread = threading.Thread(target=self.route_loop_cmb, daemon=True)
        thread.start()
    
    def route_loop_cmb(self):
        # Initialize the router socket for receiving messages
        context = zmq.Context()

        # Router socket (inbound)
        socket = context.socket(zmq.ROUTER)
        socket.bind(f"tcp://localhost:{self.port_in}")
        print(f"[ROUTER] Channel router running on port {self.port_in}")

        # Initialize the router socket for publishing messages
        pub_socket = context.socket(zmq.PUB)
        pub_socket.bind(f"tcp://localhost:{self.port_out}")
        print(f"[ROUTER]  Publish channel router running on port {self.port_out}")

        # Initialize Ack router socket
        ack_port_in = context.socket(zmq.ROUTER)
        ack_port_in.bind(f"tcp://localhost:{self.port_ack_ingress}")
        print(f"[ROUTER] Ack channel router running on port {self.port_ack_ingress}")

        # Initialize Ack publisher socket
        ack_port_out = context.socket(zmq.ROUTER)
        ack_port_out.bind(f"tcp://localhost:{self.port_ack_egress}")
        print(f"[ROUTER] Ack channel publisher running on port {self.port_ack_egress}")

        while True:
            try:
                frames = socket.recv_multipart()
                identity = frames[0]
                raw_msg = frames[-1]
                msg = CognitiveMessage.from_bytes(raw_msg)
                print(f"[Router] {self.channel_name} received message from {msg.source}")

                # ---- validation phase ---- (stub)
                # Ack message format to be determined
                if not self.validate_message(msg):
                    ack = {
                        "type": "ROUTER_ACK",
                        "status": "rejected",
                        "reason": "validation_failed",
                        "message_id": msg.message_id
                    }

                # ----Publisher Phase ----    
                for target in msg.targets:
                    pub_socket.send_multipart([
                    target.encode(),
                    msg.to_bytes()
                    ])
                    
                    print(f"[CMBRouter] Routed message from {msg.source} to {target} via {self.channel_name}")
                
                # Ack phase (after publish)
                ack = AckMessage.create(
                    msg_type="ACK",
                    ack_type="ROUTER_ACK",
                    status="SUCCESS",
                    source="CMB_ROUTER",
                    targets=[msg.source],
                    correlation_id=msg.message_id,
                    payload={
                        "type": "ROUTER_ACK",
                        "status": "published",
                        "channel": self.channel_name,
                        "message_id": msg.message_id   
                    }
                )

                ack_port_out.send_multipart([
                    identity,
                    b"",
                    AckMessage.to_bytes(ack)]
                )

                print(f"[CMBRouter] Sent ACK to {msg.source} for message {msg.message_id} identity {identity}")

            except Exception as e:
                print(f"[CMBRouter] Error routing message: {e}")
    
    def close(self):
        self.router_socket.close()
        self.context.term()
        

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\cmb_router_legecy_router_ack.py
# ==================================================

"""
Module: cmb_router.py
Location: src/core/cmb/
Version: 0.2.0

Minimal ROUTER-based channel router for the Cognitive Message Bus.

- One router per channel
- ROUTER → ROUTER forwarding
- Immediate ROUTER_ACK
- No state machine
"""

import json
from core.messages.message_module import MessageType
import zmq
import threading
import time
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.messages.ack_message import AckMessage
from src.core.cmb.cmb_channel_config import (
    get_channel_ingress_port,
    get_ack_egress_port,
    get_channel_egress_port
)


class ChannelRouter:
    def __init__(self, channel_name: str, host: str = "localhost"):
        self.channel_name = channel_name
        self.host = host

        self.router_port = get_channel_ingress_port(channel_name)
        self.module_egress_port = get_channel_egress_port(channel_name)
        self.ack_port = get_ack_egress_port(channel_name)

        self._stop_evt = threading.Event()
        self._thread = None
        print(f"[ROUTER.{self.channel_name}] router_port = {self.router_port}, ack_port = {self.ack_port}, module_egress_port = {self.module_egress_port}")

    def start(self) -> None:
        if self._thread and self._thread.is_alive():
            return
        self._stop_evt.clear()
        self._thread = threading.Thread(
            target=self._run,
            name=f"ChannelRouter[{self.channel_name}]",
            daemon=False,
        )
        self._thread.start()
        print(f"[Router.{self.channel_name}] Channel '{self.channel_name}' started")

    def stop(self) -> None:
        self._stop_evt.set()
        if self._thread:
            self._thread.join(timeout=2.0)
        print(f"[Router.{self.channel_name}] Channel '{self.channel_name}' stopped")


    def _run(self) -> None:
        ctx = zmq.Context.instance()
        print(f"[Router.{self.channel_name}] starting on {self.router_port}")
        router_sock = ctx.socket(zmq.ROUTER)
        router_sock.bind(f"tcp://{self.host}:{self.router_port}")

        module_egress_sock = ctx.socket(zmq.ROUTER)
        module_egress_sock.bind(f"tcp://{self.host}:{self.module_egress_port}")

        ack_sock = ctx.socket(zmq.ROUTER)
        ack_sock.bind(f"tcp://{self.host}:{self.ack_port}")

        poller = zmq.Poller()
        poller.register(router_sock, zmq.POLLIN)
        poller.register(module_egress_sock, zmq.POLLIN)

        print(f"[Router.{self.channel_name}] {self.channel_name} ROUTER on {self.router_port}, ")
        print(f"[Router.{self.channel_name}] {self.channel_name} ACK on {self.ack_port}")

        try:
            while not self._stop_evt.is_set():
                events = dict(poller.poll(100))

                if router_sock not in events:
                    continue
                
                # Receive message
                frames = router_sock.recv_multipart()
                sender_id = frames[0]
                payload = frames[-1]
                print(f"[Router.{self.channel_name} DEBUG] Received from sender_id={sender_id!r}")

                # ---- Parse message ----
                try:
                    obj = json.loads(payload.decode("utf-8"))
                except Exception as e:    
                    print(f"[Router.{self.channel_name} ERROR] Invalid JSON message: {e}")
                    continue

                # ---- Determine message type ----
                msg_type = obj.get("msg_type")
                print(f"[Router.{self.channel_name} DEBUG] msg_type={msg_type!r}")
                
                #---- Handle ACK messages ----
                if msg_type == "ACK":
                    ack = AckMessage.from_dict(obj)
                    print(
                        f"[Router.{self.channel_name}] {self.channel_name} "
                        f"forwarded ACK {ack.message_id} "
                        f"to {ack.targets[0]}"
                    )
                    # ---- Forward ACK to destination ----
                    dest = ack.targets[0].encode("utf-8")
                    ack_sock.send_multipart([dest, b"", payload])

                    # ---- If MESSAGE_DELIVERED_ACK, acknowledge to sender ----
                    if ack.ack_type == "MESSAGE_DELIVERED_ACK":            
                        router_ack = AckMessage.create(
                            msg_type="ACK",
                            ack_type="ROUTER_ACK",
                            status="SUCCESS",
                            source="CMB_ROUTER",
                            targets=[ack.source],
                            correlation_id=ack.correlation_id,
                            payload={"channel": self.channel_name,
                                    "status": "published",
                                    "message_id": msg.message_id,},
                        )
                        ack_dest = ack.source.encode("utf-8")
                        ack_sock.send_multipart([ack_dest, b"", AckMessage.to_bytes(router_ack)])


                    continue
                else:
                    print(f"[Router.{self.channel_name} DEBUG] Processing COMMAND message")
                    try:
                        msg = CognitiveMessage.from_dict(obj)
                    except Exception as e:
                        print(f"[Router.{self.channel_name} ERROR] Invalid COMMAND message: {e}")
                        continue

                print(
                    f"[Router.{self.channel_name}] {self.channel_name} "
                    f"received {msg.msg_type} from {msg.source}"
                )

                # ---- Forward to targets ----
                for target in msg.targets:
                    print(f"[Router.{self.channel_name} DEBUG] Sending to identity={target!r}")

                    module_egress_sock.send_multipart([
                        target.encode("utf-8"),
                        b"",
                        payload,
                    ])
                    print(
                        f"[Router.{self.channel_name}] Forwarded {msg.message_id} "
                        f"to {target}"
                    )

                # ---- Immediate ROUTER_ACK ----
                ack = AckMessage.create(
                    msg_type="ACK",
                    ack_type="ROUTER_ACK",
                    status="SUCCESS",
                    source="CMB_ROUTER",
                    targets=[msg.source],
                    correlation_id=msg.message_id,
                    payload={
                        "channel": self.channel_name,
                        "status": "published",
                        "message_id": msg.message_id,
                    },
                )

                ack_sock.send_multipart([
                    sender_id,
                    b"",
                    AckMessage.to_bytes(ack),
                ])

                print(
                    f"[Router.{self.channel_name}] ACK sent to {msg.source} "
                    f"for {msg.message_id}"
                )
        finally:
            router_sock.close()
            ack_sock.close()
            ctx.term()
            print(f"[Router.{self.channel_name}] {self.channel_name} shutdown complete")

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\endpoint_config.py
# ==================================================

from dataclasses import dataclass
from typing import Dict, Iterable, Optional

import zmq
from src.core.cmb.channel_registry import ChannelRegistry, ChannelConfig, InboundDelivery
from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_entry import LogEntry
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink

@dataclass(frozen=True)
class ChannelEndpointConfig:
    """Per-channel socket and port configuration for a module endpoint."""

    name: str
    router_port: int
    inbound_port: Optional[int]
    ack_port: Optional[int]

    # Socket types (defaults reflect your current ROUTER/DEALER design)
    outbound_socket_type: int = zmq.DEALER
    inbound_socket_type: int = zmq.DEALER
    ack_socket_type: int = zmq.DEALER

    # Delivery mode
    inbound_delivery: InboundDelivery = InboundDelivery.DIRECTED


@dataclass(frozen=True)
class MultiChannelEndpointConfig:
    """
    Configuration for a module endpoint supporting multiple channels.

    One ModuleEndpoint instance
    One module identity
    Multiple channel connections resolved from ChannelRegistry
    """

    module_id: str                        # e.g. "behavior.executor.1"
    channels: Dict[str, ChannelConfig]    # resolved ChannelConfig objects
    host: str = "localhost"
    poll_timeout_ms: int = 50

    # Logging
    log_manager = LogManager(min_severity=LogSeverity.INFO)
    log_manager.register_sink(
    FileLogSink("logs/system.jsonl")
    )

    logger = Logger("endpoint_config", log_manager)

    logger.info(
    event_type="ENDPOINT_CONFIG_INIT",
    message=f"Endpoint configuration ",
    payload={
        "note": "no payload"
    }
)

    @classmethod
    def from_channel_names(
        cls,
        *,
        module_id: str,
        channel_names: Iterable[str],
        host: str = "localhost",
        poll_timeout_ms: int = 50,
    ) -> "MultiChannelEndpointConfig":
        """
        Factory method that builds endpoint configuration
        from ChannelRegistry channel names.
        """

        channels: Dict[str, ChannelConfig] = {}
        ChannelRegistry.initialize()
        for name in channel_names:
            channels[name] = ChannelRegistry.get(name)
            
        return cls(
            module_id=module_id,
            channels=channels,
            host=host,
            poll_timeout_ms=poll_timeout_ms,
        )

    def channel_names(self) -> list[str]:
        return list(self.channels.keys())

    def get_channel(self, name: str) -> ChannelConfig:
        if name not in self.channels:
            raise KeyError(
                f"Channel '{name}' not configured for module '{self.module_id}'"
            )
        return self.channels[name]

#@dataclass(frozen=True)
#class MultiChannelEndpointConfig2:
    """All-channel configuration for a module endpoint."""

 #   module_id: str
  #  host: str = "localhost"
   # poll_timeout_ms: int = 100
    #channels: Dict[str, ChannelEndpointConfig] = field(default_factory=dict)

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\module_endpoint.py
# ==================================================

# module_endpoint.py
from __future__ import annotations

import threading
import time
import queue
from typing import Optional, Callable, Any
from typing import Dict
import zmq
import json

from src.core.cmb.utils import extract_message_id
from src.core.cmb.endpoint_config import MultiChannelEndpointConfig
from src.core.cmb.channel_registry import InboundDelivery
from src.core.cmb.transaction_registry import TransactionRegistry
from src.core.messages.ack_message import AckMessage
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_entry import LogEntry
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink


class ModuleEndpoint:
    """
    ModuleEndpoint: transport + queues boundary.

    Thread ownership model:
      - start() spins a thread
      - thread creates Context + sockets and runs poll loop
      - module logic never touches zmq sockets

    Queues:
      - _send_q: module logic -> endpoint (outbound messages)
      - _in_q: endpoint -> module logic (inbound messages)
      - _ack_q: endpoint -> module logic (ACK messages)
    """

    def __init__(
        self,
        config: MultiChannelEndpointConfig,
        *,
        logger: Optional[Callable[[str], None]] = None,
        serializer: Optional[Callable[[Any], bytes]] = None,
        deserializer: Optional[Callable[[bytes], Any]] = None,
        
    ):
        self.cfg = config

        # Logging
        self.log_manager = LogManager(min_severity=LogSeverity.INFO)
        self.log_manager.register_sink(
        FileLogSink("logs/system.jsonl")
        )
        self.logger = Logger(self.cfg.module_id, self.log_manager)

        # Old logger function fallback
        self._log = None

        # Default serializer assumes caller already provides bytes.
        self._to_bytes = serializer or (lambda x: x if isinstance(x, (bytes, bytearray)) else str(x).encode("utf-8"))
        self._from_bytes = deserializer or (lambda b: b)

        self._send_q: "queue.Queue[tuple[str,bytes, bytes]]" = queue.Queue()
        self._in_q: "queue.Queue[Any]" = queue.Queue()
        self._ack_q: "queue.Queue[Any]" = queue.Queue()

        self._stop_evt = threading.Event()
        self._thread: Optional[threading.Thread] = None

        # These exist only in endpoint thread
        self._ctx = None
        self._out_socks: dict[str, zmq.Socket] = {}
        self._in_socks: dict[str, zmq.Socket] = {}
        self._ack_socks: dict[str, zmq.Socket] = {}
        self._poller = None

        self._sock_to_channel: dict[zmq.Socket, str] = {}
        self._sock_is_ack: dict[zmq.Socket, bool] = {}

        self._tx_registry = TransactionRegistry()

    # --------------------------
    # Public API (module side)
    # --------------------------

    def start(self) -> None:
        if self._thread and self._thread.is_alive():
            return
        self._stop_evt.clear()
        self._thread = threading.Thread(target=self._run, name=f"Endpoint[{self.cfg.module_id}]", daemon=False)
        self._thread.start()
        
        self.logger.info(
            event_type="ENDPOINT_START",
            message=f"ModuleEndpoint started {self.cfg.module_id}",
            payload={
                "channels": list(self.cfg.channels.keys())
            }
        )

    def stop(self, join_timeout: float = 2.0) -> None:
        self._stop_evt.set()
        if self._thread:
            self._thread.join(timeout=join_timeout)
        
        self.logger.info(
            event_type="ENDPOINT_STOP",
            message=f"ModuleEndpoint stopped {self.cfg.module_id}",
            payload={
                "channels": list(self.cfg.channels.keys())
            }
        )

    
    def send(self, channel: str, target_id: str, payload: bytes) -> None:
        if not isinstance(payload, (bytes, bytearray)):
            raise TypeError(
                f"ModuleEndpoint.send expects bytes, got {type(payload)}"
            )
        dest = target_id.encode("utf-8")
        self._send_q.put((channel, dest, payload))


    def recv(self, timeout: Optional[float] = None) -> Optional[Any]:
        """Receive a normal inbound message (not ACK)."""
        try:
            return self._in_q.get(timeout=timeout)
        except queue.Empty:
            return None

    def recv_ack(self, timeout: Optional[float] = None) -> Optional[Any]:
        """Receive an ACK message."""
        try:
            return self._ack_q.get(timeout=timeout)
        except queue.Empty:
            return None

    def drain_incoming(self, max_items: int = 100) -> list[Any]:
        items = []
        for _ in range(max_items):
            try:
                items.append(self._in_q.get_nowait())
            except queue.Empty:
                break
        return items

    def drain_acks(self, max_items: int = 100) -> list[Any]:
        items = []
        for _ in range(max_items):
            try:
                items.append(self._ack_q.get_nowait())
            except queue.Empty:
                break
        return items

    # --------------------------
    # Endpoint thread internals
    # --------------------------

    def _run(self) -> None:
        try:
            self._setup_zmq()
            self._loop()
        except Exception as e:
            self.logger.info(
                event_type="ENDPOINT_EXCEPTION",
                message=f"ModuleEndpoint exception in {self.cfg.module_id}: {e!r}",
                payload={
                    "channels": list(self.cfg.channels.keys())
                }
            )
     
        finally:
            self._teardown_zmq()
            self.logger.info(
                    event_type="ENDPOINT_TEARDOWN",
                    message=f"ModuleEndpoint {self.cfg.module_id}  teardown complete ",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )

    def _setup_zmq(self) -> None:
        """
        Create and connect ZMQ sockets for all configured channels.
        Runs exclusively inside the endpoint thread.
        """
        self._ctx = zmq.Context.instance()

        # Poller for all inbound + ACK sockets
        self._poller = zmq.Poller()

        for ch_name, ch_cfg in self.cfg.channels.items():
            # ---------------------------
            # Outbound socket (DEALER -> ROUTER)
            # ---------------------------
            out_sock = self._ctx.socket(ch_cfg.outbound_socket_type)
            out_sock.setsockopt_string(zmq.IDENTITY, self.cfg.module_id)
            out_sock.connect(f"tcp://{self.cfg.host}:{ch_cfg.router_port}")

            self._out_socks[ch_name] = out_sock

            self.logger.info(
                event_type="ENDPOINT_OUTBOUND_SETUP",
                message=f"ModuleEndpoint {self.cfg.module_id} setup outbound {ch_name} ",
                payload={
                    "channels": list(self.cfg.channels.keys())
                }
            )

            # ---------------------------
            # Inbound socket (optional)
            # ---------------------------
            if ch_cfg.inbound_port is not None:
                if ch_cfg.inbound_delivery == InboundDelivery.BROADCAST:
                    in_sock = self._ctx.socket(zmq.SUB)
                    in_sock.setsockopt(zmq.SUBSCRIBE, b"")
                else:
                    in_sock = self._ctx.socket(zmq.DEALER)
                    in_sock.setsockopt_string(zmq.IDENTITY, self.cfg.module_id)


                # Identity is required for DEALER, ignored for SUB
                in_sock.setsockopt_string(zmq.IDENTITY, self.cfg.module_id)

                # SUB sockets must subscribe explicitly
                if ch_cfg.inbound_delivery.name == "BROADCAST":
                    in_sock.setsockopt(zmq.SUBSCRIBE, b"")

                in_sock.connect(f"tcp://{self.cfg.host}:{ch_cfg.inbound_port}")

                self._in_socks[ch_name] = in_sock
                self._sock_to_channel[in_sock] = ch_name
                self._sock_is_ack[in_sock] = False
                self._poller.register(in_sock, zmq.POLLIN)

                self.logger.info(
                    event_type="ENDPOINT_INBOUND_SETUP",
                    message=f"ModuleEndpoint {self.cfg.module_id} setup inbound {ch_name} ",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )


            # ---------------------------
            # ACK socket (optional, DIRECTED only)
            # ---------------------------
            if ch_cfg.ack_port is not None:
                ack_sock = self._ctx.socket(ch_cfg.ack_socket_type)
                ack_sock.setsockopt_string(zmq.IDENTITY, self.cfg.module_id)
                ack_sock.connect(f"tcp://{self.cfg.host}:{ch_cfg.ack_port}")

                self._ack_socks[ch_name] = ack_sock
                self._sock_to_channel[ack_sock] = ch_name
                self._sock_is_ack[ack_sock] = True
                self._poller.register(ack_sock, zmq.POLLIN)

                self.logger.info(
                    event_type="ENDPOINT_ACK_SETUP",
                    message=f"ModuleEndpoint {self.cfg.module_id} setup ACK {ch_name} ",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )


    def _teardown_zmq(self) -> None:
        # Close sockets created in thread
        for sock_dict in (self._out_socks, self._in_socks, self._ack_socks):
            if not sock_dict:
                continue
            for sock in sock_dict.values():
                try:
                    sock.close(linger=0)
                except Exception:
                    pass


        self._out_socks = None
        self._in_socks = None
        self._ack_socks = None
        self._poller = None

        # Do NOT terminate Context.instance() here; other endpoints may use it.
        self._ctx = None

    def _loop(self) -> None:
        """
        Main endpoint event loop.
        Handles outbound flushing and inbound/ACK dispatch
        across all configured channels.
        """
        while not self._stop_evt.is_set():
            """
            self._tx_registry.tick()
            self._tx_registry.cleanup_completed()
            """

            # 1) Flush outbound messages (fair, bounded)
            self._flush_outbound(max_per_tick=50)

            # 2) Poll inbound + ACK sockets
            if self._poller is None:
                time.sleep(0.01)
                continue

            try:
                events = dict(self._poller.poll(self.cfg.poll_timeout_ms))
            except zmq.ZMQError as e: 
                # Context terminated or shutting down

                self.logger.info(
                    event_type="ENDPOINT_ZMQ_ERROR",
                    message=f"ModuleEndpoint {self.cfg.module_id} poller error : {e!r}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )

                return

            # 3) Dispatch ready sockets
            for sock in events:
                is_ack = self._sock_is_ack.get(sock, False)
                self._handle_inbound(sock, is_ack=is_ack)


    def _flush_outbound(self, max_per_tick: int) -> None:
        """
        Flush outbound messages across all channels.
        Respects backpressure and preserves message ordering per channel.
        """
        sent = 0

        while sent < max_per_tick:
            try:
                ch_name, dest, payload = self._send_q.get_nowait()

            except queue.Empty:
                return

            # Get outbound socket for channel
            out_sock = self._out_socks.get(ch_name)
            if out_sock is None:

                self.logger.info(
                    event_type="ENDPOINT_NO_OUTBOUND_SOCKET",
                    message=f"ModuleEndpoint {self.cfg.module_id} out_sock None {ch_name} ",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )

                continue

            # Send message
            try:                
                message_id = extract_message_id(payload)
                tx = self._tx_registry.create(
                    message_id=message_id,
                    channel=ch_name,
                    source=self.cfg.module_id,
                    target=dest.decode("utf-8"),
                    payload=payload,
                )

                
                self.logger.info(
                    event_type="ENDPOINT_TRANSACTION_CREATED",
                    message=f"ModuleEndpoint {self.cfg.module_id} channel: {ch_name} outgoing message_id={message_id}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )

                # ROUTER addressing pattern:
                # [dest_identity][empty][payload]
                out_sock.send_multipart(
                    [payload],
                    flags=zmq.NOBLOCK
                )

                sent += 1

                self.logger.info(
                    event_type="ENDPOINT_SENT_MESSAGE",
                    message=f"ModuleEndpoint {self.cfg.module_id} sent message on channel {ch_name} to {dest!r}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )
                

            except zmq.Again:
                # Backpressure: requeue and retry next loop
                self._send_q.put((ch_name, dest, payload))
                return

    def _handle_inbound(self, sock, *, is_ack: bool) -> None:
        """
        Handles typical ROUTER->DEALER frames:
          [identity][empty?][payload]
        We keep this tolerant because your framing is still evolving.
        """
        frames = sock.recv_multipart()
        payload = frames[-1]


        if is_ack:
            ack = AckMessage.from_bytes(payload)
            event = self._tx_registry.apply_ack(ack)
            
            self.logger.info(
                    event_type="ENDPOINT_RECEIVED_ACK",
                    message=f"ModuleEndpoint {self.cfg.module_id} received ACK for correlation_id={ack.correlation_id}, event={event} ack type = {ack.ack_type}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )
            
            if event != "ERROR 1" and event != "ERROR 2":
                self._ack_q.put(ack)
            else:
                                
                self.logger.info(
                    event_type="ENDPOINT_ERROR_INVALID_ACK",
                    message=f"ModuleEndpoint {self.cfg.module_id} Received invalid ACK: {ack!r}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )

        else:
            msg_obj = CognitiveMessage.from_bytes(payload)
            self._in_q.put(msg_obj)
            message_id = msg_obj.message_id
            tx = self._tx_registry.create(
                    message_id=message_id,
                    channel = None,
                    source=msg_obj.source,
                    target=msg_obj.targets,
                    payload=payload,
                )
            
            self.logger.info(
                    event_type="ENDPOINT_CREATED_TRANSACTION",
                    message=f"ModuleEndpoint {self.cfg.module_id}  Created transaction for incoming message_id={message_id}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
            )
            
            # send ACK back
            try:
                ack = AckMessage.create(
                    msg_type="ACK",
                    ack_type="MESSAGE_DELIVERED_ACK",
                    status="SUCCESS",
                    source=self.cfg.module_id,
                    targets=[msg_obj.source],
                    correlation_id=msg_obj.message_id,
                    payload={ 
                        "status": "published",
                        "message_id": msg_obj.message_id
                    }
                )

                self.send("CC", msg_obj.source, ack.to_bytes())

                self.logger.info(
                    event_type="ENDPOINT_SENT_ACK",
                    message=f"ModuleEndpoint {self.cfg.module_id} sent ACK to {msg_obj.source}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )

            except Exception as e:
                self.logger.info(
                    event_type="ENDPOINT_ACK_SEND_ERROR",
                    message=f"ModuleEndpoint {self.cfg.module_id} outbound ACK send error: {e!r}",
                    payload={
                        "channels": list(self.cfg.channels.keys())
                    }
                )

            


# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\transaction_record.py
# ==================================================

from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
import time

from src.core.cmb.transport_state_machine import AckStateMachine, AckTransitionEvent


@dataclass
class TransactionRecord:
    """
    Tracks the lifecycle of a single outbound message exchange on the CMB.

    One TransactionRecord exists per message_id and owns:
    - The ACK state machine
    - Transition history
    - Timing and diagnostics
    """

    # -------------------------------------------------
    # Identity
    # -------------------------------------------------
    message_id: str
    event_id: str
    channel: str
    source: str
    target: str

    # -------------------------------------------------
    # Message payload (opaque to CMB)
    # -------------------------------------------------
    payload: bytes

    # -------------------------------------------------
    # Reliability core
    # -------------------------------------------------
    ack_sm: AckStateMachine = field(init=False)

    # -------------------------------------------------
    # Timeline
    # -------------------------------------------------
    created_at: float = field(default_factory=time.monotonic)
    completed_at: Optional[float] = None

    # -------------------------------------------------
    # History
    # -------------------------------------------------
    transitions: List[AckTransitionEvent] = field(default_factory=list)

    # -------------------------------------------------
    # Terminal diagnostics
    # -------------------------------------------------
    final_state: Optional[str] = None
    failure_reason: Optional[str] = None

    # -------------------------------------------------
    # Initialization
    # -------------------------------------------------
    def __post_init__(self):
        self.ack_sm = AckStateMachine(message_id=self.message_id)

    # -------------------------------------------------
    # State transition handling
    # -------------------------------------------------
    def record_transition(self, event: AckTransitionEvent) -> None:
        """
        Records a state transition produced by the ACK state machine
        and updates terminal metadata if applicable.
        """
        self.transitions.append(event)

        if self.ack_sm.is_terminal():
            self.completed_at = event.timestamp
            self.final_state = event.new_state

            if event.new_state in (
                "TIMEOUT",
                "COMPLETED_FAILURE",
                "CANCELLED",
            ):
                self.failure_reason = event.reason

    # -------------------------------------------------
    # Status helpers
    # -------------------------------------------------
    def is_complete(self) -> bool:
        return self.ack_sm.is_terminal()

    def duration(self) -> Optional[float]:
        if self.completed_at is None:
            return None
        return self.completed_at - self.created_at

    # -------------------------------------------------
    # Introspection / logging / GUI / DB
    # -------------------------------------------------
    def snapshot(self) -> Dict[str, Any]:
        """
        Lightweight, serializable view for GUI, logging, or persistence.
        """
        return {
            "message_id": self.message_id,
            "channel": self.channel,
            "source": self.source,
            "target": self.target,
            "state": self.ack_sm.state.name,
            "retry_count": self.ack_sm.retry_count,
            "created_at": self.created_at,
            "completed_at": self.completed_at,
            "duration": self.duration(),
            "final_state": self.final_state,
            "failure_reason": self.failure_reason,
        }

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\transaction_registry.py
# ==================================================

import threading
import time
from typing import Dict, Optional, Iterable

from src.core.cmb.cmb_exceptions import TransportError
from src.core.cmb.transaction_record import TransactionRecord
from src.core.cmb.transport_state_machine import AckTransitionEvent
from src.core.messages.ack_message import AckMessage
from src.core.messages.cognitive_message import CognitiveMessage


class TransactionRegistry:
    """
    Central registry for all in-flight and completed CMB transactions.

    Responsibilities:
      - Own all TransactionRecord instances
      - Route ACK events to the correct transaction
      - Drive timeout / retry ticks
      - Provide introspection and cleanup hooks
    """

    def __init__(self):
        self._lock = threading.RLock()
        self._transactions: Dict[str, TransactionRecord] = {}

    # -------------------------------------------------
    # Creation / lookup
    # -------------------------------------------------
    def create(
        self,
        *,
        message_id: str,
        channel: str,
        source: str,
        target: str,
        payload: bytes,
    ) -> TransactionRecord:
        """
        Create and register a new transaction.
        """
        with self._lock:
            if message_id in self._transactions:
                raise ValueError(f"Duplicate transaction for message_id={message_id}")

            tx = TransactionRecord(
                message_id=message_id,
                event_id=message_id,  # For now, event_id == message_id
                channel=channel,
                source=source,
                target=target,
                payload=payload,
            )

            # Register a transaction for this message_id
            self._transactions[message_id] = tx

            # Initial SEND transition
            event = tx.ack_sm.on_send()
            tx.record_transition(event)

            return tx

    def get(self, message_id: str) -> Optional[TransactionRecord]:
        with self._lock:
            return self._transactions.get(message_id)

    # -------------------------------------------------
    # ACK dispatch
    # -------------------------------------------------
    def apply_ack(self, ack: AckMessage) -> Optional[AckTransitionEvent]:
        """
        Apply an ACK message to the corresponding transaction.

        Returns the resulting AckTransitionEvent (if any).
        """
        with self._lock:
            tx = self._transactions.get(ack.correlation_id)
            if tx is None:
                # Unknown or already cleaned-up transaction
                raise TransportError("ERROR 1")

            if ack.ack_type == "ROUTER_ACK":
                event = tx.ack_sm.on_router_ack()

            elif ack.ack_type == "MESSAGE_DELIVERED_ACK":
                event = tx.ack_sm.on_msg_delivered_ack()               
            else:
                # Unknown ACK type → ignore safely
                event = "ERROR 2"
                return event

            tx.record_transition(event)
            return event

    def apply_msg_received(self, msg: CognitiveMessage) -> Optional[AckTransitionEvent]:
        """
        Apply a MSG_RECEIVED event to the corresponding transaction.

        Returns the resulting AckTransitionEvent (if any).
        """
        with self._lock:
            tx = self._transactions.get(msg.message_id)
            if tx is None:
                # Unknown or already cleaned-up transaction
                return None

            event = tx.ack_sm.on_msg_received()
            tx.record_transition(event)
            return event
    # -------------------------------------------------
    # Time-based processing
    # -------------------------------------------------
    def tick(self) -> Iterable[AckTransitionEvent]:
        """
        Drive time-based transitions (timeouts, retries).
        Should be called periodically by ModuleEndpoint.
        """
        events = []

        with self._lock:
            for tx in self._transactions.values():
                if tx.is_complete():
                    continue

                event = tx.ack_sm.tick()
                if event is not None:
                    tx.record_transition(event)
                    events.append(event)

        return events

    # -------------------------------------------------
    # Cancellation
    # -------------------------------------------------
    def cancel(self, message_id: str, reason: str = "CANCELLED") -> None:
        with self._lock:
            tx = self._transactions.get(message_id)
            if not tx or tx.is_complete():
                return

            event = tx.ack_sm.cancel(reason=reason)
            tx.record_transition(event)

    # -------------------------------------------------
    # Cleanup
    # -------------------------------------------------
    def cleanup_completed(self, max_age_sec: float = 60.0) -> None:
        """
        Remove completed transactions older than max_age_sec.
        """
        now = time.monotonic()

        with self._lock:
            to_delete = [
                mid for mid, tx in self._transactions.items()
                if tx.completed_at is not None
                and (now - tx.completed_at) > max_age_sec
            ]

            for mid in to_delete:
                del self._transactions[mid]

    # -------------------------------------------------
    # Introspection
    # -------------------------------------------------
    def snapshot(self) -> Dict[str, dict]:
        """
        Snapshot all current transactions (for GUI / debugging).
        """
        with self._lock:
            return {
                mid: tx.snapshot()
                for mid, tx in self._transactions.items()
            }

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\transport_state_machine.py
# ==================================================

from enum import Enum, auto
from dataclasses import dataclass
from typing import Optional, Any
import time

from core.messages.ack_message import AckMessage

class AckState(Enum):
    SEND_PENDING = auto()
    AWAIT_ROUTER_ACK = auto()
    AWAIT_MESSAGE_DELIVERED_ACK = auto()
    COMPLETED = auto()
    TIMEOUT = auto()
    ERROR = auto()
    CANCELLED = auto()

class AckDecision(Enum):
    NOOP = auto()
    RETRY = auto()
    COMPLETE = auto()
    FAIL = auto()

@dataclass(frozen=True)
class AckTransitionEvent:
    message_id: str
    old_state: str
    new_state: str
    reason: str
    timestamp: float
    retry_count: int
    details: Optional[Any] = None

class AckStateMachine:
    """
    Pure logic ACK state machine for a single outbound message.

    - No I/O
    - No logging
    - No threading
    - Emits AckTransitionEvent on every transition
    """

    def __init__(
        self,
        message_id: str,
        *,
        require_exec_ack: bool = True,
        allow_progress_ack: bool = True,
        router_timeout_s: float = 1.0,
        exec_timeout_s: float = 5.0,
        max_retries: int = 3,
    ):
        self.message_id = message_id

        # Policy
        self.require_exec_ack = require_exec_ack
        self.allow_progress_ack = allow_progress_ack
        self.router_timeout_s = router_timeout_s
        self.exec_timeout_s = exec_timeout_s
        self.max_retries = max_retries

        # State
        self.state = AckState.SEND_PENDING
        self.retry_count = 0

        now = time.monotonic()
        self.created_at = now
        self.last_transition_at = now

        self.router_deadline: Optional[float] = None
        self.exec_deadline: Optional[float] = None

    def _transition(
        self,
        new_state: AckState,
        *,
        reason: str,
        details: Optional[Any] = None,
    ) -> AckTransitionEvent:
        now = time.monotonic()

        event = AckTransitionEvent(
            message_id=self.message_id,
            old_state=self.state.name,
            new_state=new_state.name,
            reason=reason,
            timestamp=now,
            retry_count=self.retry_count,
            details=details,
        )

        self.state = new_state
        self.last_transition_at = now
        return event
    
    def on_send(self) -> AckTransitionEvent:
        self.router_deadline = time.monotonic() + self.router_timeout_s
        self.exec_deadline = None

        return self._transition(
            AckState.AWAIT_ROUTER_ACK,
            reason="SEND",
        )

    def on_router_ack(self) -> AckTransitionEvent:
        self.router_deadline = None
        if self.require_exec_ack:
            self.exec_deadline = time.monotonic() + self.exec_timeout_s
            return self._transition(
                AckState.AWAIT_MESSAGE_DELIVERED_ACK,
                reason="ROUTER_ACK",
            )

        return self._transition(
            AckState.COMPLETED,
            reason="NO_MSG_DELIVERED_ACK_REQUIRED",
        )


    def on_msg_delivered_ack(self) -> AckTransitionEvent:
        """
        Handle EXEC ACK from the destination module.

        Valid only when waiting for execution completion.
        """

        # --- Illegal state guard ---
        if self.state != AckState.AWAIT_MESSAGE_DELIVERED_ACK:
            return self._transition(
                self.state,
                reason="ILLEGAL_MSG_DELIVERED_ACK",
            )

        # --- Normal success / failure handling ---
        else: 
            return self._transition(
            AckState.COMPLETED,
            reason="MSG_DELIVERED_ACK_SUCCESS",
        )
        

    
    def on_msg_data(self, details: Optional[Any] = None) -> AckTransitionEvent:
        if not self.allow_progress_ack:
            return self._transition(
                self.state,
                reason="PROGRESS_ACK_IGNORED",
            )

        # Stay in EXECUTING, refresh timeout
        self.exec_deadline = time.monotonic() + self.exec_timeout_s
        return self._transition(
            AckState.EXECUTING,
            reason="PROGRESS_ACK",
            details=details,
        )

    def tick(self, now: Optional[float] = None) -> Optional[AckTransitionEvent]:
        if now is None:
            now = time.monotonic()

        if self.state == AckState.AWAIT_ROUTER_ACK:
            if self.router_deadline and now >= self.router_deadline:
                return self._handle_timeout("ROUTER_TIMEOUT")

        if self.state in (AckState.AWAIT_EXEC_ACK, AckState.EXECUTING):
            if self.exec_deadline and now >= self.exec_deadline:
                return self._handle_timeout("EXEC_TIMEOUT")

        return None

    def _handle_timeout(self, reason: str) -> AckTransitionEvent:
        self.retry_count += 1

        if self.retry_count <= self.max_retries:
            self.router_deadline = time.monotonic() + self.router_timeout_s
            self.exec_deadline = None
            return self._transition(
                AckState.SEND_PENDING,
                reason=f"{reason}_RETRY",
            )

        return self._transition(
            AckState.TIMEOUT,
            reason=f"{reason}_FAIL",
        )

    def cancel(self) -> AckTransitionEvent:
        self.router_deadline = None
        self.exec_deadline = None
        return self._transition(
            AckState.CANCELLED,
            reason="CANCEL",
        )
    def is_terminal(self) -> bool:
        return self.state in (
            AckState.COMPLETED,
            AckState.ERROR,
            AckState.TIMEOUT,
            AckState.CANCELLED,
        )

    def snapshot(self) -> dict:
        return {
            "message_id": self.message_id,
            "state": self.state.name,
            "retry_count": self.retry_count,
            "created_at": self.created_at,
            "last_transition_at": self.last_transition_at,
        }

# ==================================================
# FILE: C:\dev\agi-system\src\core\cmb\utils.py
# ==================================================

import json


def extract_message_id(payload: bytes) -> str:
    """
    Extract message_id from a serialized CognitiveMessage payload.

    Raises ValueError if message_id is missing or payload is invalid.
    """
    try:
        data = json.loads(payload.decode("utf-8"))
    except Exception as e:
        raise ValueError(f"Invalid message payload (not JSON): {e}")

    message_id = data.get("message_id")
    if not message_id:
        raise ValueError("Payload missing 'message_id'")

    return message_id

# ==================================================
# FILE: C:\dev\agi-system\src\core\cognition\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\concept_space\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\executive\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\executive\executive_stub.py
# ==================================================

"""
Module: executive_stub.py
Location: src/core/executive/
Version: 0.1.0

Simulates the Executive Layer in the AGI system. Sends messages via the Control Channel (CC)
to downstream modules (e.g., behavior controller). Can be run standalone for integration testing.

Depends on: module_endpoint >= 0.1.0, cognitive_message >= 0.1.0, cmb_channel_config >= 0.1.0
"""


import time
from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.cmb.cmb_channel_config import get_channel_port


def main():
    # Get the correct output port for CC (Control Channel)
    cc_port_pub = get_channel_port("CC") + 1  # SUB/PUB port
    cc_port_push = get_channel_port("CC") + 0  # PUSH/ROUTER port

    executive = ModuleEndpoint("executive", pub_port=cc_port_pub, push_port=cc_port_push)

    try:
        print("[ExecutiveStub] Sending message to behavior controller via CC...")
        msg = CognitiveMessage.create(
            source="executive",
            targets=["behavior"],
            payload={"directive": "start_behavior", "behavior": "explore_area"},
            priority=70
        )
        executive.send(msg)
        print("[ExecutiveStub] Message sent.")

        # Optional: wait to see if any reply comes
        # print("[ExecutiveStub] Waiting for response...")
        # response = executive.receive()
        # print(f"[ExecutiveStub] Received response: {response.payload}")

        time.sleep(1)

    finally:
        executive.close()
        print("[ExecutiveStub] Shutdown.")


if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\gui\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\gui\directive_gui.py
# ==================================================

"""directive_gui.py

Minimal GUI for entering a directive and viewing the resulting plan.

- Sends DIRECTIVE_SUBMIT to NLP via CMB
- Receives PLAN_READY from Executive
- Periodically tails the JSONL log file

This is intentionally simple (Tkinter) and can be replaced later.
"""

from __future__ import annotations

import json
import os
import threading
import time
import tkinter as tk
from tkinter import scrolledtext
import queue

from src.core.cmb.endpoint_config import MultiChannelEndpointConfig
from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.messages.cognitive_message import CognitiveMessage


class DirectiveGUI:
    def __init__(self, *, logfile: str = "logs/system.jsonl"):
        self.module_id = "GUI"
        self.logfile = logfile
        
        self.root = tk.Tk()
        self.root.title("AGI-System Demo – Directive → Plan")
        self.root.geometry("1000x700")

        self.create_widgets()

        # Decide which channels the GUI participates in.
        # Start minimal for the demo: choose the channel(s) you use in the dropdown.
        _channels = ["CC", "SMC", "VB", "BFC", "DAC", "EIG", "PC", "MC", "IC", "TC"]

        cfg = MultiChannelEndpointConfig.from_channel_names(
            module_id=self.module_id,  # Identity of this module on the bus
            channel_names=_channels,
            host="localhost",
            poll_timeout_ms=50,
        )
        self.ep = ModuleEndpoint(
            config=cfg,
            logger=lambda s: None,
            serializer=lambda msg: msg.to_bytes(),
            deserializer=lambda b: b,
        )


        self.gui_inbox = queue.Queue()
        # Poll loops
        self.root.after(100, self.poll_incoming)
        self.root.after(1500, self.poll_log)

        self.ep.start()

        self._listener_thread = threading.Thread(
            target=self._endpoint_listener,
            daemon=True
        )
        self._listener_thread.start() 


    def create_widgets(self):
        # Directive input
        tk.Label(self.root, text="Directive").pack(anchor="w")
        self.directive_text = scrolledtext.ScrolledText(self.root, height=6)
        self.directive_text.pack(fill="x", padx=6, pady=4)

        # Preferred output format (optional)
        fmt_frame = tk.Frame(self.root)
        fmt_frame.pack(fill="x", padx=6)
        tk.Label(fmt_frame, text="Preferred output format (optional)").pack(side="left")
        self.format_entry = tk.Entry(fmt_frame)
        self.format_entry.pack(side="left", fill="x", expand=True, padx=6)

        # Buttons
        btn_frame = tk.Frame(self.root)
        btn_frame.pack(fill="x", padx=6, pady=6)
        tk.Button(btn_frame, text="Send directive", command=self.on_send).pack(side="left")
        tk.Button(btn_frame, text="Clear", command=self.on_clear).pack(side="left", padx=6)

        # Output display
        tk.Label(self.root, text="System output").pack(anchor="w")
        self.output_text = scrolledtext.ScrolledText(self.root, height=16)
        self.output_text.pack(fill="both", expand=True, padx=6, pady=4)

        # Log tail
        tk.Label(self.root, text="System log (tail)").pack(anchor="w")
        self.log_text = scrolledtext.ScrolledText(self.root, height=10)
        self.log_text.pack(fill="both", expand=False, padx=6, pady=4)

        self._last_log_size = 0
    
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def _endpoint_listener(self):

        while True:
            try:
               msg = self.ep.recv(timeout=1.0)
               if msg is not None:
                   self.gui_inbox.put(msg)
            except Exception:
                pass

    def on_send(self):
        directive = self.directive_text.get("1.0", "end").strip()
        preferred = self.format_entry.get().strip()

        if not directive:
            return

        msg = CognitiveMessage.create(
            schema_version=str(CognitiveMessage.get_schema_version()),
            msg_type="DIRECTIVE_SUBMIT",
            msg_version="0.1.0",
            source=self.module_id,
            targets=["NLP"],
            context_tag=None,
            correlation_id=None,
            payload={
                "directive_text": directive,
                "context": {
                    "preferred_output_format": preferred or None,
                    "ui_timestamp": time.time(),
                    "llm_model_id": "LLM_STUB",
                },
            },
            priority=60,
            ttl=60.0,
            signature="",
        )

        self.ep.send("CC", "NLP", msg.to_bytes())
        self._append_output(f"[GUI] Sent DIRECTIVE_SUBMIT to NLP (message_id={msg.message_id})\n")

    def on_clear(self):
        self.directive_text.delete("1.0", "end")
        self.output_text.delete("1.0", "end")

    def _append_output(self, s: str):
        self.output_text.insert("end", s)
        self.output_text.see("end")

    def poll_incoming(self):
        try:
            while True:
                msg = self.gui_inbox.get_nowait()

                if isinstance(msg, CognitiveMessage):
                    if msg.msg_type == "PLAN_READY":
                        self._append_output("\n=== PLAN_READY ===\n")
                        self._append_output(json.dumps(msg.payload, indent=2) + "\n")
                    elif msg.msg_type == "CLARIFICATION_REQUEST":
                        self._append_output("\n=== CLARIFICATION_REQUEST ===\n")
                        self._append_output(json.dumps(msg.payload, indent=2) + "\n")
                    else:
                        self._append_output(f"\n=== {msg.msg_type} ===\n")
                        self._append_output(json.dumps(msg.payload, indent=2) + "\n")
                else:
                    break
        except queue.Empty:
            pass

        self.root.after(100, self.poll_incoming)


    def poll_log(self):
        if not self.root.winfo_exists():
            return

        try:
            if not os.path.exists(self.logfile):
                self.root.after(500, self.poll_log)
                return

            size = os.path.getsize(self.logfile)
            if size == self._last_log_size:
                self.root.after(500, self.poll_log)
                return

            # Read new content (simple approach for demo)
            with open(self.logfile, "r", encoding="utf-8") as f:
                lines = f.readlines()[-50:]  # tail 50

            self.log_text.delete("1.0", "end")
            self.log_text.insert("end", "".join(lines))
            self.log_text.see("end")
            self._last_log_size = size

        finally:
            self.root.after(1500, self.poll_log)

    def on_close(self):
        try:
            self.ep.stop()
        finally:
            self.root.destroy()

    def run(self):
        self.root.mainloop()


def main():
    gui = DirectiveGUI(logfile="logs/system.jsonl")
    gui.run()


if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\intent_extractor.py
# ==================================================

from __future__ import annotations

from dataclasses import dataclass
from src.core.intent.interfaces import IntentExtractionInterface
from src.core.intent.llm_adapter_base import LLMAdapter
from src.core.intent.schema import from_dict, IntentValidationError
from src.core.intent.models import IntentObject


@dataclass
class IntentExtractor(IntentExtractionInterface):
    llm_adapter: LLMAdapter
    min_confidence: float = 0.85

    def extract_intent(self, directive_text: str) -> IntentObject:
        raw = self.llm_adapter.classify_directive(directive_text)

        intent = from_dict(raw)  # validates schema/enums/ranges

        # Confidence gating
        if intent.confidence_score < self.min_confidence:
            # Mark as clarification-required; keep immutable by returning a new instance.
            return IntentObject(
                intent_id=intent.intent_id,
                directive_source=intent.directive_source,
                directive_type=intent.directive_type,
                planning_required=intent.planning_required,
                urgency_level=intent.urgency_level,
                risk_level=intent.risk_level,
                expected_response_type=intent.expected_response_type,
                confidence_score=intent.confidence_score,
                domain_context=intent.domain_context,
                suggested_modules=intent.suggested_modules,
                execution_constraints=intent.execution_constraints,
                clarification_required=True,
            )

        return intent

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\interfaces.py
# ==================================================

from __future__ import annotations

from abc import ABC, abstractmethod
from src.core.intent.models import IntentObject


class IntentExtractionInterface(ABC):
    @abstractmethod
    def extract_intent(self, directive_text: str) -> IntentObject:
        """
        Accepts a directive string and returns a populated IntentObject.
        """
        raise NotImplementedError

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\llm_adapter_base.py
# ==================================================

from __future__ import annotations

from abc import ABC, abstractmethod
from typing import Any


class LLMAdapter(ABC):
    @abstractmethod
    def classify_directive(self, directive_text: str) -> dict[str, Any]:
        """
        Returns a JSON-like dict conforming to the IntentObject schema.
        The extractor will validate/convert it into an IntentObject.
        """
        raise NotImplementedError

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\llm_adapter_mock.py
# ==================================================

from __future__ import annotations

from typing import Any
import uuid

from src.core.intent.llm_adapter_base import LLMAdapter


class MockLLMAdapter(LLMAdapter):
    """
    Lightweight deterministic mock used for development and tests.
    This lets you validate architecture flow without any real LLM.
    """

    def classify_directive(self, directive_text: str) -> dict[str, Any]:
        text = (directive_text or "").strip().lower()

        # Very simple heuristic simulation of an LLM classification
        if any(k in text for k in ["explain", "history", "what is", "summarize", "define"]):
            dtype = "cognitive"
            planning = False
            expected = "textual_response"
            confidence = 0.90
        elif any(k in text for k in ["compare", "evaluate", "trade-off", "analyze", "recommend"]):
            dtype = "analytical"
            planning = False
            expected = "textual_response"
            confidence = 0.85
        elif any(k in text for k in ["build", "implement", "create", "design", "produce", "generate a plan"]):
            dtype = "goal_oriented"
            planning = True
            expected = "plan"
            confidence = 0.80
        elif any(k in text for k in ["move", "actuate", "drive", "turn on", "turn off", "robot"]):
            dtype = "behavioral"
            planning = True
            expected = "action"
            confidence = 0.80
        elif any(k in text for k in ["monitor", "watch", "alert", "notify", "detect"]):
            dtype = "supervisory"
            planning = True
            expected = "monitoring_process"
            confidence = 0.80
        else:
            dtype = "cognitive"
            planning = False
            expected = "textual_response"
            confidence = 0.60

        return {
            "intent_id": str(uuid.uuid4()),
            "directive_source": "human",
            "directive_type": dtype,
            "planning_required": planning,
            "urgency_level": "normal",
            "risk_level": "none",
            "expected_response_type": expected,
            "confidence_score": confidence,
            "domain_context": None,
            "suggested_modules": ["nlp", "knowledge_store"] if not planning else ["planner", "executive"],
            "execution_constraints": None,
            "clarification_required": False,
        }

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\logging.py
# ==================================================

from __future__ import annotations

import json
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from src.core.intent.models import IntentObject


@dataclass
class IntentLogger:
    """
    Simple JSONL logger. Works on low-power machines and is easy to replay later.
    """
    log_path: Path

    def log(self, directive_text: str, intent: IntentObject, route: str, extra: dict[str, Any] | None = None) -> None:
        record = {
            "timestamp": time.time(),
            "directive_text": directive_text,
            "intent": intent.to_dict(),
            "route": route,
            "extra": extra or {},
        }
        self.log_path.parent.mkdir(parents=True, exist_ok=True)
        with self.log_path.open("a", encoding="utf-8") as f:
            f.write(json.dumps(record) + "\n")

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\models.py
# ==================================================

from __future__ import annotations

from dataclasses import dataclass, field, asdict
from enum import Enum
from typing import Any, Optional
import uuid


class DirectiveSource(str, Enum):
    HUMAN = "human"
    PERCEPTION = "perception"
    INTERNAL = "internal"


class DirectiveType(str, Enum):
    COGNITIVE = "cognitive"
    ANALYTICAL = "analytical"
    GOAL_ORIENTED = "goal_oriented"
    BEHAVIORAL = "behavioral"
    SUPERVISORY = "supervisory"


class UrgencyLevel(str, Enum):
    LOW = "low"
    NORMAL = "normal"
    HIGH = "high"
    CRITICAL = "critical"


class RiskLevel(str, Enum):
    NONE = "none"
    LOW = "low"
    MODERATE = "moderate"
    HIGH = "high"


class ExpectedResponseType(str, Enum):
    TEXTUAL_RESPONSE = "textual_response"
    STRUCTURED_DATA = "structured_data"
    PLAN = "plan"
    ACTION = "action"
    MONITORING_PROCESS = "monitoring_process"


@dataclass(frozen=True)
class IntentObject:
    """
    Immutable intent representation produced by the intent extractor.
    Downstream control logic should depend on these fields, not on raw text.
    """
    intent_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    directive_source: DirectiveSource = DirectiveSource.HUMAN
    directive_type: DirectiveType = DirectiveType.COGNITIVE
    planning_required: bool = False
    urgency_level: UrgencyLevel = UrgencyLevel.NORMAL
    risk_level: RiskLevel = RiskLevel.NONE
    expected_response_type: ExpectedResponseType = ExpectedResponseType.TEXTUAL_RESPONSE
    confidence_score: float = 0.5

    # Optional / extended fields
    domain_context: Optional[str] = None
    suggested_modules: tuple[str, ...] = ()
    execution_constraints: Optional[dict[str, Any]] = None
    clarification_required: bool = False

    def to_dict(self) -> dict[str, Any]:
        d = asdict(self)
        # Enums serialize as values
        d["directive_source"] = self.directive_source.value
        d["directive_type"] = self.directive_type.value
        d["urgency_level"] = self.urgency_level.value
        d["risk_level"] = self.risk_level.value
        d["expected_response_type"] = self.expected_response_type.value
        return d

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\router.py
# ==================================================

from __future__ import annotations

from dataclasses import dataclass
from .models import IntentObject


class Route(str):
    DIRECT_RESPONSE = "direct_response"
    INVOKE_PLANNER = "invoke_planner"
    REQUEST_CLARIFICATION = "request_clarification"


@dataclass
class DirectiveRouter:
    """
    Deterministic routing based on the IntentObject only.
    """
    risk_requires_confirmation: bool = False  # reserved for later policy hooks

    def route(self, intent: IntentObject) -> str:
        if intent.clarification_required:
            return Route.REQUEST_CLARIFICATION

        if intent.planning_required:
            return Route.INVOKE_PLANNER

        return Route.DIRECT_RESPONSE

# ==================================================
# FILE: C:\dev\agi-system\src\core\intent\schema.py
# ==================================================

from __future__ import annotations

from typing import Any
from src.core.intent.models import (
    IntentObject,
    DirectiveSource,
    DirectiveType,
    UrgencyLevel,
    RiskLevel,
    ExpectedResponseType,
)


class IntentValidationError(ValueError):
    pass


def validate_confidence(score: float) -> None:
    if not isinstance(score, (float, int)):
        raise IntentValidationError("confidence_score must be a number.")
    if score < 0.0 or score > 1.0:
        raise IntentValidationError("confidence_score must be in [0.0, 1.0].")


def from_dict(data: dict[str, Any]) -> IntentObject:
    """
    Convert a dict (typically LLM JSON output) into an IntentObject,
    enforcing enum validity and basic constraints.
    """
    required = [
        "intent_id",
        "directive_source",
        "directive_type",
        "planning_required",
        "urgency_level",
        "risk_level",
        "expected_response_type",
        "confidence_score",
    ]
    missing = [k for k in required if k not in data]
    if missing:
        raise IntentValidationError(f"Missing required fields: {missing}")

    validate_confidence(float(data["confidence_score"]))

    try:
        intent = IntentObject(
            intent_id=str(data["intent_id"]),
            directive_source=DirectiveSource(str(data["directive_source"])),
            directive_type=DirectiveType(str(data["directive_type"])),
            planning_required=bool(data["planning_required"]),
            urgency_level=UrgencyLevel(str(data["urgency_level"])),
            risk_level=RiskLevel(str(data["risk_level"])),
            expected_response_type=ExpectedResponseType(str(data["expected_response_type"])),
            confidence_score=float(data["confidence_score"]),
            domain_context=data.get("domain_context"),
            suggested_modules=tuple(data.get("suggested_modules", []) or []),
            execution_constraints=data.get("execution_constraints"),
            clarification_required=bool(data.get("clarification_required", False)),
        )
    except Exception as e:
        raise IntentValidationError(f"Invalid schema or enum value: {e}") from e

    return intent

# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\execution_context.py
# ==================================================

from dataclasses import dataclass
from typing import Optional


@dataclass(frozen=True)
class ExecutionContext:
    """
    Captures the execution linkage for a log entry.

    This structure enables end-to-end tracing of work,
    events, and message transactions across the system.
    """

    work_id: str
    # Identifies the overall unit of work (human directive,
    # sensor-triggered response, autonomous goal, etc.).

    event_id: str
    # Identifies the specific event within the work lifecycle
    # that produced this log entry.

    transaction_id: Optional[str] = None
    # Identifies a message exchange or transactional sequence
    # (especially relevant for CMB tracing).

    parent_event_id: Optional[str] = None
    # Links this event to the event that caused it, enabling
    # causal graphs and replay trees.

# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\file_log_sink.py
# ==================================================

import json
from pathlib import Path
from typing import Optional
import threading

from src.core.logging.log_entry import LogEntry


class FileLogSink:
    """
    Log sink that persists log entries to an append-only JSONL file.

    Each LogEntry is written as a single JSON object per line,
    enabling efficient tailing, replay, and offline analysis.
    """

    def __init__(self, logfile_path: str):
        self._path = Path(logfile_path)
        self._lock = threading.Lock()

        # Ensure parent directory exists
        self._path.parent.mkdir(parents=True, exist_ok=True)

        # Open file in append mode, line-buffered
        self._file = open(self._path, "a", encoding="utf-8")

    def emit(self, entry: LogEntry) -> None:
        """
        Persist a log entry to disk.

        This method must not raise exceptions outward.
        """
        try:
            record = {
                "log_id": entry.log_id,
                "timestamp": entry.timestamp,
                "severity": entry.severity.name,
                "source_module": entry.source_module,
                "event_type": entry.event_type,
                "message": entry.message,
                "payload": entry.payload,
                "context": (
                    vars(entry.context)
                    if entry.context is not None
                    else None
                ),
            }
            with self._lock:
                self._file.write(json.dumps(record) + "\n")
                self._file.flush()

        except Exception:
            # Never allow logging to break the system
            pass

    def close(self) -> None:
        """
        Close the underlying file handle.
        """
        try:
            with self._lock:
                self._file.close()
        except Exception:
            pass

# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\gui_log_sink.py
# ==================================================

import queue
from typing import Optional

from src.core.logging.log_entry import LogEntry


class GuiSubscriptionSink:
    """
    Log sink that forwards log entries to a GUI via a thread-safe queue.

    This sink performs no UI operations and is safe to use
    from any thread.
    """

    def __init__(self, gui_queue: queue.Queue):
        self._queue = gui_queue

    def emit(self, entry: LogEntry) -> None:
        """
        Forward a log entry to the GUI queue.

        This method must not block or raise exceptions.
        """
        try:
            self._queue.put_nowait(entry)
        except Exception:
            pass

# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\log_entry.py
# ==================================================

from dataclasses import dataclass, field
from typing import Any, Dict, Optional
import time
import uuid

from src.core.logging.log_severity import LogSeverity
from src.core.logging.execution_context import ExecutionContext


@dataclass
class LogEntry:
    """
    Atomic record of a single observable system event.

    LogEntry is the fundamental unit used for:
    - diagnostics
    - observability
    - replay
    - learning
    - threat and error analysis
    """

    log_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    # Unique identifier for this log entry.
    # Used for indexing, correlation, and replay.

    timestamp: float = field(default_factory=time.time)
    # Wall-clock time when the event occurred.
    # High resolution is important for ordering and analysis.

    severity: LogSeverity = LogSeverity.INFO
    # Semantic importance of the event.

    source_module: str = ""
    # Logical system module that produced this entry
    # (e.g., GUI, NLP, EXEC, PLANNER, ROUTER).

    event_type: str = ""
    # Machine-readable symbolic name for the event
    # (e.g., DIRECTIVE_RECEIVED, PLAN_CREATED).

    message: Optional[str] = None
    # Human-readable summary of the event.
    # Optional but strongly recommended.

    payload: Dict[str, Any] = field(default_factory=dict)
    # Structured data associated with the event.
    # Must be JSON-serializable.
    # Never required for system correctness.

    context: Optional[ExecutionContext] = None
    # Execution linkage information tying this log entry
    # into a larger work / event / transaction sequence.

# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\log_manager.py
# ==================================================

from typing import  Protocol

from src.core.logging.log_entry import LogEntry
from src.core.logging.log_severity import LogSeverity

class LogSink(Protocol):
    """
    Abstract destination for log entries.

    A LogSink may persist logs, stream them, buffer them,
    or forward them to another subsystem.
    """

    def emit(self, entry: LogEntry) -> None:
        """
        Receive a log entry for processing.

        Must not block the caller.
        Must not raise exceptions outward.
        """

class LogManager:
    """
    Central coordinator for system logging.

    LogManager is responsible for accepting log entries
    and distributing them to registered sinks.
    """

    def __init__(self, *, min_severity: LogSeverity = LogSeverity.INFO):
        self._min_severity = min_severity
        self._sinks: list[LogSink] = []

    def register_sink(self, sink: LogSink) -> None:
        """
        Register a new sink to receive log entries.

        Sinks may include file writers, UI streams,
        replay buffers, or future analyzers.
        """
        self._sinks.append(sink)

    def log(self, entry: LogEntry) -> None:
        """
        Submit a log entry to the logging subsystem.

        This method is intended to be called by all modules.
        It must be fast and non-blocking.
        """

        if entry.severity.value < self._min_severity.value:
            return

        for sink in self._sinks:
            try:
                sink.emit(entry)
            except Exception:
                # Logging must never destabilize the system.
                # Failures here are intentionally swallowed.
                pass

class Logger:
    """
    Convenience façade bound to a specific module.
    """

    def __init__(self, module_id: str, manager: LogManager):
        self._module_id = module_id
        self._manager = manager

    def info(self, *, event_type: str, message: str, context=None, payload=None):
        self._manager.log(
            LogEntry(
                severity=LogSeverity.INFO,
                source_module=self._module_id,
                event_type=event_type,
                message=message,
                context=context,
                payload=payload or {},
            )
        )

# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\log_severity.py
# ==================================================

from enum import Enum, auto


class LogSeverity(Enum):
    """
    Semantic severity level for log entries.

    Used for filtering, alerting, analysis, and future
    error/threat detection subsystems.
    """

    TRACE = auto()      # Extremely fine-grained execution detail
    DEBUG = auto()      # Developer-focused diagnostic information
    INFO = auto()       # Normal system operation
    WARNING = auto()    # Unexpected but recoverable condition
    ERROR = auto()      # Operation failed, system continued
    CRITICAL = auto()   # System integrity or safety at risk

# ==================================================
# FILE: C:\dev\agi-system\src\core\logging\logging_wrapper.py
# ==================================================

"""logging_wrapper.py

A minimal, swappable logging wrapper.

Goals (per architecture):
- every module can log consistently
- logs can be written to file now, database later
- supports a GUI viewer by writing newline-delimited JSON
"""

from __future__ import annotations

import json
import threading
import time
from dataclasses import dataclass, asdict
from typing import Any, Optional


@dataclass
class LogEntry:
    timestamp: float               # Epoch seconds when the log entry was created
    level: str                     # e.g., DEBUG, INFO, WARN, ERROR
    module: str                    # Module emitting the log
    event: str                     # Short event name, e.g., "directive_received"
    message: str                   # Human-readable message
    data: dict[str, Any]           # Structured details for later analysis
    wid: Optional[str] = None      # Work id (if known)
    xid: Optional[str] = None      # Exchange/transaction id (if known)
    eid: Optional[str] = None      # Event id (if known)


class JsonlLogger:
    """Thread-safe JSONL logger."""

    def __init__(self, logfile_path: str):
        self._path = logfile_path
        self._lock = threading.RLock()

    def write(self, entry: LogEntry) -> None:
        line = json.dumps(asdict(entry), ensure_ascii=False)
        with self._lock:
            with open(self._path, "a", encoding="utf-8") as f:
                f.write(line + "\n")

    def log(
        self,
        *,
        level: str,
        module: str,
        event: str,
        message: str,
        data: Optional[dict[str, Any]] = None,
        wid: Optional[str] = None,
        xid: Optional[str] = None,
        eid: Optional[str] = None,
    ) -> None:
        self.write(
            LogEntry(
                timestamp=time.time(),
                level=level,
                module=module,
                event=event,
                message=message,
                data=data or {},
                wid=wid,
                xid=xid,
                eid=eid,
            )
        )

# ==================================================
# FILE: C:\dev\agi-system\src\core\memory\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\messages\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\messages\ack_message.py
# ==================================================

"""
Module: ack_message.py
Location: src/core/messages/
Version: 0.1.0

Defines the standard ack message structure used throughout the AGI system's Cognitive Message Bus (CMB).
This module should be treated as canonical and never duplicated or modified outside of version-controlled updates.
"""

import uuid
import time
import json
from dataclasses import dataclass, asdict



@dataclass
class AckMessage:
    message_id: str              # Global unique identifier
    msg_type: str                # Acknowledgement intent
    ack_type: str                # Acknowledgement type
    status: str                  # Acknowledgement status
    source: str                  # Sending module
    targets: list[str]           # Intended recipients
    correlation_id: str | None   # Request–response linkage
    payload: dict                # Message content


    
    @staticmethod
    def create(
       msg_type: str,
       ack_type: str,
       status: str,
       source: str,
       targets: list[str],
       correlation_id: str | None,
       payload: dict,

    ) -> "AckMessage":
        return AckMessage(
            message_id=str(uuid.uuid4()),
            msg_type = msg_type,
            ack_type = ack_type,
            status = status,
            source=source,
            targets=targets,
            correlation_id = correlation_id,
            payload=payload,

        )


    def to_json(self) -> str:
        return json.dumps(asdict(self))

    def to_bytes(self) -> bytes:
        return self.to_json().encode("utf-8")

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_bytes(cls, data: bytes) -> "AckMessage":
        if not isinstance(data, (bytes, bytearray)):
            raise TypeError(
                f"AckMessage.from_bytes expects bytes, got {type(data)}"
            )
        try:
            obj = json.loads(data.decode("utf-8"))
        except Exception as e:
            raise ValueError(f"Invalid ACK payload JSON: {e}") from e

        return cls.from_dict(obj)


    @staticmethod
    def from_json(json_str: str) -> "AckMessage":
        return AckMessage(**json.loads(json_str))
    
    @classmethod
    def from_dict(cls, data: dict) -> "AckMessage":
        try:
            return cls(
                message_id=data["message_id"],
                msg_type=data["msg_type"],
                ack_type=data["ack_type"],
                status=data["status"],
                source=data["source"],
                targets=data.get("targets", []),
                correlation_id=data.get("correlation_id"),
                payload=data.get("payload", {})
            )
        except KeyError as e:
            raise ValueError(f"Missing required AckMessage field: {e}")


# ==================================================
# FILE: C:\dev\agi-system\src\core\messages\cognitive_message.py
# ==================================================

"""
Module: cognitive_message.py
Location: src/core/messages/
Version: 0.1.0

Defines the standard message structure used throughout the AGI system's Cognitive Message Bus (CMB).
This module should be treated as canonical and never duplicated or modified outside of version-controlled updates.
"""

import uuid
import time
import json
from dataclasses import dataclass, asdict
from src.core.messages.ack_message import AckMessage



@dataclass
class CognitiveMessage:
    message_id: str              # Global unique identifier
    schema_version: str          # Message schema version
    msg_type: str                # Semantic intent
    msg_version: str             # Message-type version
    source: str                  # Sending module
    targets: list[str]           # Intended recipients
    context_tag: str | None      # Goal / task context
    correlation_id: str | None   # Request–response linkage
    payload: dict                # Message content
    priority: int                # 0–100
    timestamp: float             # Epoch seconds
    ttl: float                   # Time-to-live (seconds)
    signature: str | None        # Optional integrity/auth


    
    @staticmethod
    def create(
       schema_version: str,
       msg_type: str,
       msg_version: str,
       source: str,
       targets: list[str],
       context_tag: str | None,
       correlation_id: str | None,
       payload: dict,
       priority: int = 50,
       ttl: float = 10.0,
       signature: str = ""
    ) -> "CognitiveMessage":
        message_id = str(uuid.uuid4())
        correlation_id = message_id if correlation_id is None else correlation_id
        return CognitiveMessage(
            message_id=message_id,
            schema_version = CognitiveMessage.get_schema_version(),
            msg_type = msg_type,
            msg_version = msg_version,
            source=source,
            targets=targets,
            context_tag = context_tag,
            correlation_id = correlation_id,
            payload=payload,
            priority=priority,
            timestamp=time.time(),
            ttl=ttl,
            signature=signature
        )
    
    @staticmethod
    def system_ack(payload: dict) -> "AckMessage":
        return AckMessage(
           message_id=str(uuid.uuid4()),
           source="CMB_ROUTER",
           targets=[],
           payload=payload,
           priority=0
    )

    @staticmethod
    def get_schema_version():
        return 1

    def is_expired(self) -> bool:
        return (time.time() - self.timestamp) > self.ttl

    def to_json(self) -> str:
        return json.dumps(asdict(self))

    def to_bytes(self) -> bytes:
        return self.to_json().encode("utf-8")

    def to_dict(self) -> dict:
        return asdict(self)

    @staticmethod
    def from_bytes(data: bytes) -> "CognitiveMessage":
        obj = json.loads(data.decode("utf-8"))
        return CognitiveMessage(**obj)

    @staticmethod
    def from_json(json_str: str) -> "CognitiveMessage":
        return CognitiveMessage(**json.loads(json_str))

    @classmethod
    def from_dict(cls, data: dict) -> "CognitiveMessage":
        try:
            return cls(
                message_id=data["message_id"],
                schema_version=data["schema_version"],
                msg_type=data["msg_type"],
                msg_version=data.get("msg_version", "0.0.0"),
                source=data["source"],
                targets=data.get("targets", []),
                context_tag=data.get("context_tag"),
                correlation_id=data.get("correlation_id"),
                payload=data.get("payload", {}),
                priority=data.get("priority", 0),
                ttl=data.get("ttl", 0),
                timestamp=data.get("timestamp"),
                signature=data.get("signature"),
            )
        except KeyError as e:
            raise ValueError(f"Missing required CognitiveMessage field: {e}")
    

# ==================================================
# FILE: C:\dev\agi-system\src\core\messages\message_module.py
# ==================================================

# ======================================================
# message_module.py (Message Architecture Scaffold)
# ======================================================

from enum import Enum
from typing import Any, Dict, Optional
from datetime import datetime

# ------------------------------------------------------
# Standardized Message Types
# ------------------------------------------------------
class MessageType(Enum):
    LOG_ENTRY = "log_entry"
    COMMAND = "command"
    DIRECTIVE = "directive"
    INTROSPECTION = "introspection"
    VISUAL = "visual"
    AUDIO = "audio"
    LANGUAGE = "language"
    SENSOR = "sensor"
    REWARD = "reward"
    ERROR = "error"
    EXIT = "exit"
    STATUS = "status"
    QUERY = "query"
    RESPONSE = "response"
    ROUTE = "route"
    DIAGNOSTIC = "diagnostic"

# ------------------------------------------------------
# Standardized Module Names
# ------------------------------------------------------
class ModuleName(Enum):
    QUESTION_GENERATOR = "question_generator"
    AEM = "aem"
    NLP = "nlp"
    PERCEPTION = "perception"
    SELF_TALK = "self_talk"
    CMB_LOGGER = "cmb_logger"
    DIAGNOSTICS = "diagnostics"
    ROUTER = "cmb_router"
    LOGGER_UI = "logger_ui"
    GUI_DIRECTIVE = "gui_directive"

#

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\behavior_module.py
# ==================================================

"""
Behavior Stub Module

Receives messages from the Cognitive Message Bus and logs them.
Refactored to use the CommonModuleLoop.
"""

from __future__ import annotations

import time

from src.core.cmb.endpoint_config import MultiChannelEndpointConfig
from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink
from src.core.modules.common_module_loop import CommonModuleLoop


MODULE_ID = "behavior"


def main():
    # -----------------------------
    # Logging setup
    # -----------------------------
    log_manager = LogManager(min_severity=LogSeverity.INFO)
    log_manager.register_sink(FileLogSink("logs/system.jsonl"))
    logger = Logger(MODULE_ID, log_manager)

    logger.info(
        event_type="BEHAVIOR_INIT",
        message="Behavior module initializing",
    )

    # -----------------------------
    # Endpoint setup
    # -----------------------------
    channels = [
        "CC", "SMC", "VB", "BFC", "DAC",
        "EIG", "PC", "MC", "IC", "TC"
    ]

    cfg = MultiChannelEndpointConfig.from_channel_names(
        module_id=MODULE_ID,
        channel_names=channels,
        host="localhost",
        poll_timeout_ms=50,
    )

    endpoint = ModuleEndpoint(
        config=cfg,
        logger=None,  # logging handled separately
        serializer=lambda msg: msg.to_bytes(),
        deserializer=lambda b: b,
    )

    endpoint.start()

    # -----------------------------
    # Message handler
    # -----------------------------
    def handle_message(msg):
        logger.info(
            event_type="BEHAVIOR_MESSAGE_RECEIVED",
            message="Behavior received message",
            payload={
                "msg_type": msg.msg_type,
                "source": msg.source,
                "message_id": msg.message_id,
            },
        )

        # For now: no response logic
        # Future: behavior selection, sequencing, execution

    # -----------------------------
    # Optional lifecycle hooks
    # -----------------------------
    def on_start():
        logger.info(
            event_type="BEHAVIOR_START",
            message="Behavior module started",
        )

    def on_shutdown():
        logger.info(
            event_type="BEHAVIOR_SHUTDOWN",
            message="Behavior module shutting down",
        )

    # -----------------------------
    # Run loop
    # -----------------------------
    loop = CommonModuleLoop(
        module_id=MODULE_ID,
        endpoint=endpoint,
        logger=logger,
        on_message=handle_message,
        on_start=on_start,
        on_shutdown=on_shutdown,
    )

    try:
        loop.start()
    except KeyboardInterrupt:
        logger.info(
            event_type="BEHAVIOR_INTERRUPT",
            message="Behavior interrupted by user",
        )
        loop.stop()


if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\common.py
# ==================================================

"""common.py

Shared utilities for module stubs.
"""

from __future__ import annotations

from src.core.cmb.cmb_channel_config import (
    get_channel_ingress_port,
    get_channel_egress_port,
    get_ack_egress_port,
)
from src.core.cmb.endpoint_config import ChannelEndpointConfig, MultiChannelEndpointConfig2


def default_endpoint_config(module_id: str, *, host: str = "localhost") -> MultiChannelEndpointConfig2:
    """Create a minimal MultiChannelEndpointConfig for the CC channel."""
    ch_name = "CC"
    ch_cfg = ChannelEndpointConfig(
        name=ch_name,
        router_port=get_channel_ingress_port(ch_name),
        inbound_port=get_channel_egress_port(ch_name),
        ack_port=get_ack_egress_port(ch_name),
    )
    return MultiChannelEndpointConfig2(module_id=module_id, host=host, channels={ch_name: ch_cfg})

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\common_module_loop.py
# ==================================================

import time
import threading
from typing import Optional, Callable

from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.logging.log_manager import Logger


class CommonModuleLoop:
    """
    Generic execution loop for all CMB-connected modules.

    This loop mediates between ModuleEndPoint and module-specific logic.
    """

    def __init__(
        self,
        *,
        module_id: str,
        endpoint: ModuleEndpoint,
        logger: Logger,
        on_message: Callable,
        on_start: Optional[Callable] = None,
        on_tick: Optional[Callable] = None,
        on_shutdown: Optional[Callable] = None,
        poll_interval: float = 0.1,
    ):
        self.module_id = module_id
        self.endpoint = endpoint
        self.logger = logger

        self.on_message = on_message
        self.on_start = on_start
        self.on_tick = on_tick
        self.on_shutdown = on_shutdown

        self.poll_interval = poll_interval
        self._stop_evt = threading.Event()

    def start(self):
        self.logger.info(
            event_type="MODULE_LOOP_START",
            message="Module loop starting",
            payload={"module_id": self.module_id},
        )

        if self.on_start:
            try:
                self.on_start()
            except Exception as e:
                self.logger.info(
                    event_type="MODULE_START_ERROR",
                    message=str(e),
                )

        self.run()

    def stop(self):
        self.logger.info(
            event_type="MODULE_LOOP_STOP_REQUEST",
            message="Stop requested",
        )
        self._stop_evt.set()

    def run(self):
        try:
            while not self._stop_evt.is_set():
                # Receive message (non-blocking or short timeout)
                msg = self.endpoint.recv(timeout=self.poll_interval)

                if msg is not None:
                    self.logger.info(
                        event_type="MODULE_MESSAGE_RECV",
                        message="Message received",
                        payload={
                            "msg_type": msg.msg_type,
                            "source": msg.source,
                            "message_id": msg.message_id,
                        },
                    )

                    try:
                        self.on_message(msg)
                    except Exception as e:
                        self.logger.info(
                            event_type="MODULE_MESSAGE_HANDLER_ERROR",
                            message="Exception in module message handler",
                            payload={
                                "exception_type": type(e).__name__,
                                "exception": str(e),
                            },
                        )   

                # Optional periodic work
                if self.on_tick:
                    try:
                        self.on_tick()
                    except Exception as e:
                        self.logger.info(
                            event_type="MODULE_TICK_ERROR",
                            message=str(e),
                        )

        finally:
            if self.on_shutdown:
                try:
                    self.on_shutdown()
                except Exception as e:
                    self.logger.info(
                        event_type="MODULE_SHUTDOWN_ERROR",
                        message=str(e),
                    )

            self.endpoint.stop()

            self.logger.info(
                event_type="MODULE_LOOP_EXIT",
                message="Module loop exited",
            )

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\executive_module.py
# ==================================================

"""
Executive Module (Stub)

Receives plans from the Planner and produces a task queue for execution.
For now:
- Accept PLAN_READY
- Create a deterministic task queue (stub)
- Send TASK_QUEUE_READY to GUI (so you can see end-to-end flow)
"""

from __future__ import annotations

import time
import uuid
from typing import Any, Dict, List

from src.core.cmb.endpoint_config import MultiChannelEndpointConfig
from src.core.cmb.module_endpoint import ModuleEndpoint

from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink

from src.core.modules.common_module_loop import CommonModuleLoop
from src.core.messages.cognitive_message import CognitiveMessage


MODULE_ID = "EXEC"


def _make_task_queue_from_plan(plan: Dict[str, Any]) -> Dict[str, Any]:
    """
    Convert a plan dict into a simple task queue structure.

    This is intentionally deterministic and minimal for the demo.
    """
    now = time.time()
    plan_id = plan.get("plan_id") or str(uuid.uuid4())
    steps = plan.get("steps") or []

    tasks: List[Dict[str, Any]] = []

    # If planner sent steps, create one task per step
    if isinstance(steps, list) and steps:
        for idx, step in enumerate(steps, start=1):
            tasks.append(
                {
                    "task_id": str(uuid.uuid4()),
                    "task_index": idx,
                    "name": step.get("description") or f"Plan Step {idx}",
                    "assigned_module": step.get("assigned_module") or "behavior",
                    "created_at": now,
                    "status": "QUEUED",
                }
            )
    else:
        # Fallback: always create at least one task
        tasks.append(
            {
                "task_id": str(uuid.uuid4()),
                "task_index": 1,
                "name": "Execute plan (stub)",
                "assigned_module": "behavior",
                "created_at": now,
                "status": "QUEUED",
            }
        )

    return {
        "queue_id": str(uuid.uuid4()),
        "created_at": now,
        "plan_id": plan_id,
        "task_count": len(tasks),
        "tasks": tasks,
        "status": "READY",
    }


def _create_message(
    *,
    msg_type: str,
    source: str,
    targets: list[str],
    payload: dict,
    correlation_id: str | None,
    context_tag: str | None,
    priority: int = 50,
    ttl: float = 60.0,
) -> CognitiveMessage:
    """
    Safe wrapper around CognitiveMessage.create() using the full signature
    your current code expects.
    """
    return CognitiveMessage.create(
        schema_version=str(CognitiveMessage.get_schema_version()),
        msg_type=msg_type,
        msg_version="0.1.0",
        source=source,
        targets=targets,
        context_tag=context_tag,
        correlation_id=correlation_id,
        payload=payload,
        priority=priority,
        ttl=ttl,
        signature="",
    )


def main():
    # -----------------------------
    # Logging
    # -----------------------------
    log_manager = LogManager(min_severity=LogSeverity.INFO)
    log_manager.register_sink(FileLogSink("logs/system.jsonl"))
    logger = Logger(MODULE_ID, log_manager)

    logger.info(event_type="EXEC_INIT", message="Executive module initializing")

    # -----------------------------
    # Endpoint
    # -----------------------------
    channels = ["CC", "SMC", "VB", "BFC", "DAC", "EIG", "PC", "MC", "IC", "TC"]

    cfg = MultiChannelEndpointConfig.from_channel_names(
        module_id=MODULE_ID,
        channel_names=channels,
        host="localhost",
        poll_timeout_ms=50,
    )

    endpoint = ModuleEndpoint(
        config=cfg,
        logger=None,  # we use LogManager/Logger instead
        serializer=lambda msg: msg.to_bytes(),
        deserializer=lambda b: b,
    )

    # -----------------------------
    # Handler
    # -----------------------------
    def handle_message(msg: CognitiveMessage):
        if msg.msg_type != "PLAN_READY":
            return

        logger.info(
            event_type="EXEC_PLAN_RECEIVED",
            message="PLAN_READY received from Planner",
            payload={
                "source": msg.source,
                "message_id": msg.message_id,
                "correlation_id": msg.correlation_id,
            },
        )

        plan = msg.payload.get("plan") or {}
        task_queue = _make_task_queue_from_plan(plan)

        logger.info(
            event_type="EXEC_TASK_QUEUE_CREATED",
            message="Task queue created from plan",
            payload={
                "plan_id": task_queue.get("plan_id"),
                "queue_id": task_queue.get("queue_id"),
                "task_count": task_queue.get("task_count"),
            },
        )

        # Send to GUI so the demo shows end-to-end flow
        out = _create_message(
            msg_type="TASK_QUEUE_READY",
            source=MODULE_ID,
            targets=["GUI"],
            payload={
                "task_queue": task_queue,
                "plan": plan,
            },
            correlation_id=msg.message_id,
            context_tag=plan.get("plan_id"),
            priority=60,
            ttl=60.0,
        )

        endpoint.send("CC", "GUI", out.to_bytes())

        logger.info(
            event_type="EXEC_TASK_QUEUE_EMITTED",
            message="TASK_QUEUE_READY sent to GUI",
            payload={
                "target": "GUI",
                "queue_id": task_queue.get("queue_id"),
                "correlation_id": msg.message_id,
            },
        )

    # -----------------------------
    # Lifecycle hooks
    # -----------------------------
    def on_start():
        endpoint.start()
        logger.info(event_type="EXEC_START", message="Executive module started")

    def on_shutdown():
        logger.info(event_type="EXEC_SHUTDOWN", message="Executive module shutting down")

    # -----------------------------
    # Run loop
    # -----------------------------
    loop = CommonModuleLoop(
        module_id=MODULE_ID,
        endpoint=endpoint,
        logger=logger,
        on_message=handle_message,
        on_start=on_start,
        on_shutdown=on_shutdown,
    )

    try:
        loop.start()
    except KeyboardInterrupt:
        logger.info(event_type="EXEC_INTERRUPT", message="Executive interrupted by user")
        loop.stop()


if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\executive_module_legacy.py
# ==================================================

"""executive_module.py

Stub Executive module.

Responsibilities:
- Receive DirectiveDerivative from NLP
- Decide whether to initiate work (stubbed as always true unless clarifications needed)
- Create WorkInstance
- Request plan from Planner
- Receive Plan and forward to GUI for display
- (Later) enqueue tasks for execution
"""

from __future__ import annotations

import time
import uuid
from dataclasses import asdict
from typing import Any, Dict

from core.cmb.channel_registry import ChannelRegistry
from core.cmb.endpoint_config import MultiChannelEndpointConfig

from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.logging.logging_wrapper import JsonlLogger

from src.core.architecture.agi_system_dataclasses import WorkInstance, WorkStatus, Plan, TaskSpec


def _create_work_from_derivative(derivative: Dict[str, Any]) -> WorkInstance:
    """Create a WorkInstance from a DirectiveDerivative dict."""
    wid = str(uuid.uuid4())
    now = time.time()
    raw = str(derivative.get("raw_directive", ""))
    return WorkInstance(
        wid=wid,
        title=(raw[:80] + ("..." if len(raw) > 80 else "")),
        goal=raw,
        status=WorkStatus.CREATED,
        created_at=now,
        updated_at=now,
        context=dict(derivative.get("context", {})),
        origin="human_directive",
        tags=["directive"],
    )


def run_executive_module(*, logfile: str = "logs/system.jsonl") -> None:
    module_id = "EXEC"
    logger = JsonlLogger(logfile)

    ChannelRegistry.initialize()
    
    # Decide which channels the GUI participates in.
    # Start minimal for the demo: choose the channel(s) you use in the dropdown.
    _channels = ["CC", "SMC", "VB", "BFC", "DAC", "EIG", "PC", "MC", "IC", "TC"]

    cfg = MultiChannelEndpointConfig.from_channel_names(
        module_id=module_id,  # Identity of this module on the bus
        channel_names=_channels,
        host="localhost",
        poll_timeout_ms=50,
    )

    ep = ModuleEndpoint(
        config=cfg,
        logger=lambda s: logger.log(level="DEBUG", module=module_id, event="endpoint", message=s),
        serializer=lambda msg: msg.to_bytes(),
        deserializer=lambda b: b,
    )
    ep.start()

    logger.log(level="INFO", module=module_id, event="started", message="Executive module started")

    # Minimal in-memory cache for current work (demo)
    current_work: dict[str, WorkInstance] = {}

    try:
        while True:
            msg = ep.recv(timeout=0.2)
            if msg is None:
                continue

            if not isinstance(msg, CognitiveMessage):
                continue

            # 1) DirectiveDerivative -> create WorkInstance -> send to Planner
            if msg.msg_type == "DIRECTIVE_DERIVATIVE":
                derivative = dict(msg.payload.get("derivative", {}))

                # If NLP produced clarification questions, we pause and ask GUI
                clarifications = derivative.get("system_clarification_questions", []) or []
                if clarifications:
                    resp = CognitiveMessage.create(
                        schema_version=str(CognitiveMessage.get_schema_version()),
                        msg_type="CLARIFICATION_REQUEST",
                        msg_version="0.1.0",
                        source=module_id,
                        targets=["GUI"],
                        context_tag=None,
                        correlation_id=msg.correlation_id,
                        payload={"questions": clarifications},
                        priority=60,
                        ttl=60.0,
                        signature="",
                    )
                    ep.send("CC", "GUI", resp.to_bytes())
                    logger.log(level="INFO", module=module_id, event="clarification_requested", message="Asked GUI for clarification", data={"questions": clarifications})
                    continue

                work = _create_work_from_derivative(derivative)
                work.status = WorkStatus.PLANNING
                work.updated_at = time.time()
                current_work[work.wid] = work

                logger.log(
                    level="INFO",
                    module=module_id,
                    event="work_created",
                    message="WorkInstance created from directive",
                    data={"wid": work.wid, "correlation_id": msg.correlation_id},
                )

                plan_req = CognitiveMessage.create(
                    schema_version=str(CognitiveMessage.get_schema_version()),
                    msg_type="PLAN_REQUEST",
                    msg_version="0.1.0",
                    source=module_id,
                    targets=["PLANNER"],
                    context_tag=work.wid,
                    correlation_id=msg.correlation_id,
                    payload={
                        "wid": work.wid,
                        "work": asdict(work),
                        "derivative": derivative,
                    },
                    priority=55,
                    ttl=60.0,
                    signature="",
                )

                ep.send("CC", "PLANNER", plan_req.to_bytes())
                logger.log(level="INFO", module=module_id, event="plan_requested", message="Requested plan from Planner", data={"wid": work.wid})
                continue

            # 2) Plan -> forward to GUI and mark work as planned
            if msg.msg_type == "PLAN_RESPONSE":
                wid = str(msg.payload.get("wid") or "")
                plan_dict = dict(msg.payload.get("plan") or {})

                work = current_work.get(wid)
                if work:
                    work.status = WorkStatus.PLANNED
                    work.updated_at = time.time()

                logger.log(level="INFO", module=module_id, event="plan_received", message="Received plan from Planner", data={"wid": wid, "task_count": len(plan_dict.get("tasks", []))})

                gui_msg = CognitiveMessage.create(
                    schema_version=str(CognitiveMessage.get_schema_version()),
                    msg_type="PLAN_READY",
                    msg_version="0.1.0",
                    source=module_id,
                    targets=["GUI"],
                    context_tag=wid,
                    correlation_id=msg.correlation_id,
                    payload={"wid": wid, "plan": plan_dict},
                    priority=50,
                    ttl=60.0,
                    signature="",
                )

                ep.send("CC", "GUI", gui_msg.to_bytes())
                continue

    except KeyboardInterrupt:
        logger.log(level="INFO", module=module_id, event="stopped", message="Executive module stopped")
    finally:
        ep.stop()


if __name__ == "__main__":
    run_executive_module()

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\nlp_module.py
# ==================================================

"""
NLP Module

Receives directives from the GUI, normalizes them,
and forwards structured directives to the Planner.
"""

from __future__ import annotations

import time

from src.core.cmb.endpoint_config import MultiChannelEndpointConfig
from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink
from src.core.modules.common_module_loop import CommonModuleLoop
from src.core.messages.cognitive_message import CognitiveMessage


MODULE_ID = "NLP"


def main():
    # -----------------------------
    # Logging setup
    # -----------------------------
    log_manager = LogManager(min_severity=LogSeverity.INFO)
    log_manager.register_sink(FileLogSink("logs/system.jsonl"))
    logger = Logger(MODULE_ID, log_manager)

    logger.info(
        event_type="NLP_INIT",
        message="NLP module initializing",
    )

    # -----------------------------
    # Endpoint setup
    # -----------------------------
    channels = [
        "CC", "SMC", "VB", "BFC", "DAC",
        "EIG", "PC", "MC", "IC", "TC"
    ]

    cfg = MultiChannelEndpointConfig.from_channel_names(
        module_id=MODULE_ID,
        channel_names=channels,
        host="localhost",
        poll_timeout_ms=50,
    )

    endpoint = ModuleEndpoint(
        config=cfg,
        logger=logger.info,
        serializer=lambda msg: msg.to_bytes(),
        deserializer=lambda b: b,
    )

    # -----------------------------
    # Message handler
    # -----------------------------
    def handle_message(msg):
        if msg.msg_type != "DIRECTIVE_SUBMIT":
            return

        logger.info(
            event_type="NLP_DIRECTIVE_RECEIVED",
            message="Directive received",
            payload={
                "source": msg.source,
                "message_id": msg.message_id,
            },
        )

        directive_text = msg.payload.get("directive_text")

        normalized_payload = {
            "original_text": directive_text,
            "received_at": time.time(),
            "source": msg.source,
        }

        out_msg = CognitiveMessage.create(
            schema_version=str(CognitiveMessage.get_schema_version()),
            msg_type="DIRECTIVE_NORMALIZED",
            msg_version="0.1.0",
            source=MODULE_ID,
            targets=["PLANNER"],
            context_tag=None,
            correlation_id=msg.message_id,
            payload=normalized_payload,
        )

        endpoint.send("CC", "PLANNER", out_msg.to_bytes())

        logger.info(
            event_type="NLP_DIRECTIVE_EMITTED",
            message="Normalized directive sent to planner",
            payload={
                "target": "PLANNER",
                "correlation_id": msg.message_id,
            },
        )

    # -----------------------------
    # Lifecycle hooks
    # -----------------------------
    def on_start():
        endpoint.start()
        logger.info(
            event_type="NLP_START",
            message="NLP module started",
        )

    def on_shutdown():
        logger.info(
            event_type="NLP_SHUTDOWN",
            message="NLP module shutting down",
        )

    # -----------------------------
    # Run loop
    # -----------------------------
    loop = CommonModuleLoop(
        module_id=MODULE_ID,
        endpoint=endpoint,
        logger=logger,
        on_message=handle_message,
        on_start=on_start,
        on_shutdown=on_shutdown,
    )

    try:
        loop.start()
    except KeyboardInterrupt:
        logger.info(
            event_type="NLP_INTERRUPT",
            message="NLP interrupted by user",
        )
        loop.stop()


if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\planner_module.py
# ==================================================

"""
Planner Module

Receives normalized directives from NLP and produces
a preliminary execution plan for the Executive.
"""

from __future__ import annotations

import time
import uuid

from src.core.cmb.endpoint_config import MultiChannelEndpointConfig
from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.logging.log_manager import LogManager, Logger
from src.core.logging.log_severity import LogSeverity
from src.core.logging.file_log_sink import FileLogSink
from src.core.modules.common_module_loop import CommonModuleLoop
from src.core.messages.cognitive_message import CognitiveMessage


MODULE_ID = "PLANNER"


def main():
    # -----------------------------
    # Logging setup
    # -----------------------------
    log_manager = LogManager(min_severity=LogSeverity.INFO)
    log_manager.register_sink(FileLogSink("logs/system.jsonl"))
    logger = Logger(MODULE_ID, log_manager)

    logger.info(
        event_type="PLANNER_INIT",
        message="Planner module initializing",
    )

    # -----------------------------
    # Endpoint setup
    # -----------------------------
    channels = [
        "CC", "SMC", "VB", "BFC", "DAC",
        "EIG", "PC", "MC", "IC", "TC"
    ]

    cfg = MultiChannelEndpointConfig.from_channel_names(
        module_id=MODULE_ID,
        channel_names=channels,
        host="localhost",
        poll_timeout_ms=50,
    )

    endpoint = ModuleEndpoint(
        config=cfg,
        logger=None,
        serializer=lambda msg: msg.to_bytes(),
        deserializer=lambda b: b,
    )

    # -----------------------------
    # Message handler
    # -----------------------------
    def handle_message(msg):
        if msg.msg_type != "DIRECTIVE_NORMALIZED":
            return

        logger.info(
            event_type="PLANNER_DIRECTIVE_RECEIVED",
            message="Normalized directive received from NLP",
            payload={
                "source": msg.source,
                "message_id": msg.message_id,
                "correlation_id": msg.correlation_id,
            },
        )

        directive = msg.payload.get("original_text")

        # -----------------------------
        # Stub plan generation
        # -----------------------------
        plan_id = str(uuid.uuid4())

        plan = {
            "plan_id": plan_id,
            "created_at": time.time(),
            "source_directive": directive,
            "steps": [
                {
                    "step_id": "step-1",
                    "description": "Analyze directive",
                    "assigned_module": "EXECUTIVE",
                }
            ],
            "status": "READY",
        }

        out_msg = CognitiveMessage.create(
            schema_version=str(CognitiveMessage.get_schema_version()),
            msg_type="PLAN_READY",
            msg_version="0.1.0",
            source=MODULE_ID,
            targets=["EXEC"],
            context_tag=None,
            payload={
                "plan": plan,
            },
            correlation_id=msg.correlation_id,
        )

        endpoint.send("CC", "EXEC", out_msg.to_bytes())

        logger.info(
            event_type="PLANNER_PLAN_EMITTED",
            message="Plan sent to Executive",
            payload={
                "plan_id": plan_id,
                "target": "EXEC",
                "correlation_id": msg.message_id,
            },
        )

    # -----------------------------
    # Lifecycle hooks
    # -----------------------------
    def on_start():
        endpoint.start()
        logger.info(
            event_type="PLANNER_START",
            message="Planner module started",
        )

    def on_shutdown():
        logger.info(
            event_type="PLANNER_SHUTDOWN",
            message="Planner module shutting down",
        )

    # -----------------------------
    # Run loop
    # -----------------------------
    loop = CommonModuleLoop(
        module_id=MODULE_ID,
        endpoint=endpoint,
        logger=logger,
        on_message=handle_message,
        on_start=on_start,
        on_shutdown=on_shutdown,
    )

    try:
        loop.start()
    except KeyboardInterrupt:
        logger.info(
            event_type="PLANNER_INTERRUPT",
            message="Planner interrupted by user",
        )
        loop.stop()


if __name__ == "__main__":
    main()

# ==================================================
# FILE: C:\dev\agi-system\src\core\modules\planner_module_legacy.py
# ==================================================

"""planner_module.py

Stub Planner module.

Responsibilities:
- Receive PLAN_REQUEST
- Produce a Plan (list of TaskSpecs) suitable for Executive scheduling

This is a placeholder: it creates a small deterministic plan derived from the directive.
"""

from __future__ import annotations

import time
import uuid
from dataclasses import asdict
from typing import Any, Dict, List

from core.cmb.channel_registry import ChannelRegistry
from core.cmb.endpoint_config import MultiChannelEndpointConfig

from src.core.cmb.module_endpoint import ModuleEndpoint
from src.core.messages.cognitive_message import CognitiveMessage
from src.core.logging.logging_wrapper import JsonlLogger

from src.core.architecture.agi_system_dataclasses import Plan, TaskSpec


def _build_plan(*, wid: str, work: Dict[str, Any], derivative: Dict[str, Any]) -> Plan:
    raw = str(derivative.get("raw_directive", ""))
    now = time.time()

    tasks: List[TaskSpec] = [
        TaskSpec(
            task_id=str(uuid.uuid4()),
            name="Interpret directive",
            description="Normalize the directive into actionable intent and constraints.",
            inputs={"raw_directive": raw},
            outputs={"intent": "(stub)"},
            depends_on=[],
            created_at=now,
        ),
        TaskSpec(
            task_id=str(uuid.uuid4()),
            name="Draft response outline",
            description="Create a preliminary outline of the deliverable.",
            inputs={"intent": "(stub)"},
            outputs={"outline": "(stub)"},
            depends_on=[],
            created_at=now,
        ),
        TaskSpec(
            task_id=str(uuid.uuid4()),
            name="Finalize plan",
            description="Assemble final task plan and confirm deliverable format.",
            inputs={"preferred_output_format": derivative.get("preferred_output_format")},
            outputs={"plan": "(stub)"},
            depends_on=[],
            created_at=now,
        ),
    ]

    return Plan(
        plan_id=str(uuid.uuid4()),
        wid=wid,
        created_at=now,
        tasks=tasks,
        rationale="Stub planner: generates a simple 3-step plan for all directives.",
    )


def run_planner_module(*, logfile: str = "logs/system.jsonl") -> None:
    module_id = "PLANNER"
    logger = JsonlLogger(logfile)

    ChannelRegistry.initialize()
    
    # Decide which channels the PLANNER participates in.
    # Start minimal for the demo: choose the channel(s) you use in the dropdown.
    _channels = ["CC", "SMC", "VB", "BFC", "DAC", "EIG", "PC", "MC", "IC", "TC"]

    cfg = MultiChannelEndpointConfig.from_channel_names(
        module_id=module_id,  # Identity of this module on the bus
        channel_names=_channels,
        host="localhost",
        poll_timeout_ms=50,
    )
    ep = ModuleEndpoint(
        config=cfg,
        logger=lambda s: logger.log(level="DEBUG", module=module_id, event="endpoint", message=s),
        serializer=lambda msg: msg.to_bytes(),
        deserializer=lambda b: b,
    )
    ep.start()

    logger.log(level="INFO", module=module_id, event="started", message="Planner module started")

    try:
        while True:
            msg = ep.recv(timeout=0.2)
            if msg is None:
                continue

            if not isinstance(msg, CognitiveMessage):
                continue

            if msg.msg_type != "PLAN_REQUEST":
                continue

            wid = str(msg.payload.get("wid") or "")
            work = dict(msg.payload.get("work") or {})
            derivative = dict(msg.payload.get("derivative") or {})

            logger.log(level="INFO", module=module_id, event="plan_request_received", message="PLAN_REQUEST received", data={"wid": wid})

            plan = _build_plan(wid=wid, work=work, derivative=derivative)

            resp = CognitiveMessage.create(
                schema_version=str(CognitiveMessage.get_schema_version()),
                msg_type="PLAN_RESPONSE",
                msg_version="0.1.0",
                source=module_id,
                targets=["EXEC"],
                context_tag=wid,
                correlation_id=msg.correlation_id,
                payload={"wid": wid, "plan": asdict(plan)},
                priority=50,
                ttl=60.0,
                signature="",
            )

            ep.send("CC", "EXEC", resp.to_bytes())
            logger.log(level="INFO", module=module_id, event="plan_sent", message="Plan sent to Executive", data={"wid": wid, "task_count": len(plan.tasks)})

    except KeyboardInterrupt:
        logger.log(level="INFO", module=module_id, event="stopped", message="Planner module stopped")
    finally:
        ep.stop()


if __name__ == "__main__":
    run_planner_module()

# ==================================================
# FILE: C:\dev\agi-system\src\core\monitoring\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\perception\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\perception\perception_module.py
# ==================================================

# ==============================================
# perception_module.py
# ==============================================

import time
from multiprocessing import Queue
from modules.cmb_router import CMBRouter
from modules.message_module import send_message, build_message

def perception_loop(router_inbox, module_queue, module_name):
    queue = module_queue

    send_message(router_inbox, {
        "target": "cmb_logger",
        "module": module_name,
        "message": "[Perception] Perception module active."
    })

    counter = 0
    running = True

    while running:
        try:
            msg = queue.get(timeout=0.1)
            if isinstance(msg, dict):
                if msg.get("target", "").lower() in [module_name.lower(), "all"]:
                    if msg.get("content", "").lower() == "exit":
                        send_message(router_inbox, {
                            "target": "cmb_logger",
                            "module": module_name,
                            "message": "[Perception] Shutdown signal received. Exiting..."
                        })
                        running = False
                        break
        except:
            pass

        # Simulate visual frame input
        sample_input = {
            "source": "camera",
            "type": "visual",
            "content": f"frame_{counter}"
        }

        send_message(router_inbox, {
            "target": "cmb_logger",
            "module": module_name,
            "message": sample_input
        })

        time.sleep(2)
        counter += 1

# ==================================================
# FILE: C:\dev\agi-system\src\core\planning\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\questioning\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\reasoning\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\reflection\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\core\utils\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\factory_qc\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\factory_qc\adapters\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\factory_qc\cli\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\factory_qc\config\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\factory_qc\domain\__init__.py
# ==================================================


# ==================================================
# FILE: C:\dev\agi-system\src\factory_qc\services\__init__.py
# ==================================================

