from docx import Document
from pathlib import Path
from datetime import datetime


def create_docx(title: str, sections: list[str]) -> dict:
    """
    Create a Word document with a title and section headings.
    Returns metadata about the created artifact.
    """

    artifacts_dir = Path("artifacts")
    artifacts_dir.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{title.replace(' ', '_')}_{timestamp}.docx"
    filepath = artifacts_dir / filename

    doc = Document()
    doc.add_heading(title, level=1)

    for section in sections:
        doc.add_heading(section, level=2)
        doc.add_paragraph("")

    doc.save(filepath)

    return {
        "artifact_type": "docx",
        "path": str(filepath),
        "title": title,
        "sections": sections
    }
